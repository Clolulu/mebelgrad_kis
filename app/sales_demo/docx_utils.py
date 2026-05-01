import io
from datetime import datetime

from docx import Document
from flask import send_file


def _safe(value):
    return str(value or "")


def _add_company_block(doc, company):
    if not company:
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    head = table.rows[0].cells
    head[0].text = "Реквизит"
    head[1].text = "Значение"

    rows = [
        ("Компания", getattr(company, "company_name", "")),
        ("Краткое название", getattr(company, "short_name", "")),
        ("ИНН/КПП", f"{getattr(company, 'inn', '')} / {getattr(company, 'kpp', '')}".strip(" /")),
        ("Адрес", getattr(company, "legal_address", "")),
        ("Телефон", getattr(company, "phone", "")),
        ("Email", getattr(company, "email", "")),
        ("Руководитель", getattr(company, "ceo", "")),
    ]
    for label, value in rows:
        if value:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = _safe(value)
    doc.add_paragraph()


def create_sales_doc(title, company, order=None, customer=None, items=None, notes=None):
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph()
    _add_company_block(doc, company)

    if customer:
        doc.add_heading("Клиент", level=1)
        doc.add_paragraph(f"ФИО/Наименование: {_safe(getattr(customer, 'name', ''))}")
        doc.add_paragraph(f"Телефон: {_safe(getattr(customer, 'phone', ''))}")
        doc.add_paragraph(f"Email: {_safe(getattr(customer, 'email', ''))}")
        if getattr(customer, "registration_address", None):
            doc.add_paragraph(f"Адрес: {_safe(customer.registration_address)}")
        doc.add_paragraph()

    if order:
        doc.add_heading("Данные заказа", level=1)
        doc.add_paragraph(f"Номер заказа: {_safe(getattr(order, 'order_number', ''))}")
        doc.add_paragraph(f"Дата: {getattr(order, 'created_at', datetime.now()).strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph(f"Статус: {_safe(getattr(order, 'status', ''))}")
        doc.add_paragraph(f"Сумма: {_safe(getattr(order, 'total_amount', 0))} руб.")
        doc.add_paragraph(f"Адрес доставки: {_safe(getattr(order, 'delivery_address', ''))}")
        doc.add_paragraph()

    if items:
        doc.add_heading("Состав заказа", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        head = table.rows[0].cells
        head[0].text = "Товар"
        head[1].text = "Кол-во"
        head[2].text = "Цена"
        head[3].text = "Сумма"
        head[4].text = "Наличие"
        for item in items:
            row = table.add_row().cells
            qty = int(item.get("quantity", 0))
            price = float(item.get("unit_price", 0))
            row[0].text = _safe(item.get("name"))
            row[1].text = str(qty)
            row[2].text = f"{price:.2f}"
            row[3].text = f"{qty * price:.2f}"
            row[4].text = _safe(item.get("stock"))
        doc.add_paragraph()

    if notes:
        doc.add_paragraph(notes)

    doc.add_paragraph("Подпись продавца: ____________________")
    doc.add_paragraph("Подпись клиента: _____________________")
    return doc


def duplicate_contract_doc(company, customer, items, delivery_address):
    doc = create_sales_doc(
        "Договор купли-продажи (2 экземпляра)",
        company,
        customer=customer,
        items=items,
        notes=f"Адрес доставки: {_safe(delivery_address)}",
    )
    doc.add_page_break()
    doc.add_heading("Экземпляр №2", level=1)
    doc.add_paragraph("Содержимое экземпляра идентично экземпляру №1.")
    return doc


def save_doc(doc, filename):
    payload = io.BytesIO()
    doc.save(payload)
    payload.seek(0)
    return send_file(
        payload,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )
