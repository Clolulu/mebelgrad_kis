import io
import json
import os
import re
import urllib.request
from collections import defaultdict
from datetime import datetime, timedelta
from functools import wraps
from types import SimpleNamespace

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from flask import flash, redirect, render_template, request, send_file, url_for
from flask_login import current_user, login_required
from sqlalchemy import func, or_

from app.finance import finance_bp
from app.models import (
    BalanceSnapshot,
    BudgetItem,
    CashCalendarItem,
    CompanyProfile,
    Customer,
    IndirectExpense,
    InventoryBatch,
    Payment,
    PlanFactDeviation,
    PurchaseOrder,
    PurchaseOrderItem,
    SalesOrder,
    SalesOrderItem,
    Stock,
    Supplier,
    db,
)


def get_company_profile():
    company = CompanyProfile.query.first()
    if company:
        return company
    return SimpleNamespace(
        company_name="Мебельград",
        short_name="Мебельград",
        legal_form="ООО",
        inn="1234567890",
        kpp="123456789",
        ogrn="1234567890123",
        okved="31.01",
        tax_system="ОСН",
        employees_count=50,
        legal_address="г. Москва, ул. Примерная, д. 1",
        actual_address="г. Москва, ул. Примерная, д. 1",
        phone="+7 (495) 123-45-67",
        email="info@mebelgrad.ru",
        website="www.mebelgrad.ru",
        bank_name="Сбербанк",
        bank_bik="044525225",
        correspondent_account="30101810400000000225",
        settlement_account="40702810000000001234",
        ceo="Иванов Иван Иванович",
        ceo_position="Генеральный директор",
        ceo_signature_url="/static/images/signature.png",
        signature_url="/static/images/signature.png",
        chief_accountant_name="Петрова Петрова Петровна",
        chief_accountant_signature_url="/static/images/signature.png",
        seal_url="/static/images/seal.png",
        logo_url="/static/images/logo.png",
        print_footer="Отчет сформирован автоматически",
    )


def resolve_local_static_paths(html):
    static_root = os.path.abspath(
        os.path.join(os.path.dirname(__file__), "..", "..", "static")
    )
    static_root = static_root.replace("\\", "/")

    def replace_src(match):
        attr, quote, path = match.groups()
        return f'{attr}={quote}file:///{static_root}/{path}{quote}'

    html = re.sub(
        r'(src|href)=(["\'])/static/([^"\']+)\2',
        replace_src,
        html,
        flags=re.IGNORECASE,
    )
    html = re.sub(
        r'url\((["\']?)/static/([^"\')]+)(["\']?)\)',
        lambda m: f'url({m.group(1)}file:///{static_root}/{m.group(2)}{m.group(3)})',
        html,
        flags=re.IGNORECASE,
    )
    return html


def format_currency(value):
    try:
        return f"{value:,.2f}".replace(",", " ").replace(".", ",")
    except Exception:
        return str(value)


def format_percent(value):
    try:
        return f"{value:.2f}%".replace('.', ',')
    except Exception:
        return str(value)


def format_value(value):
    if isinstance(value, bool):
        return "Да" if value else "Нет"
    if isinstance(value, (int, float)):
        return format_currency(value)
    return str(value)


def create_docx_document(title, company=None, period=None, snapshot_date=None, now=None):
    doc = Document()

    default_logo_url = (
        "https://encrypted-tbn0.gstatic.com/images?q=tbn:ANd9GcQqjBb3N-5K12RjxM-PFGwrzRDeXoYRWvH9Ww&s"
    )
    logo_bytes = None
    logo_path = None

    logo_url = None
    if company and getattr(company, 'logo_url', None):
        logo_url = company.logo_url
    if not logo_url:
        logo_url = default_logo_url

    if logo_url.startswith('http'):
        try:
            request_obj = urllib.request.Request(
                logo_url,
                headers={
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
                                  '(KHTML, like Gecko) Chrome/115.0 Safari/537.36'
                },
            )
            with urllib.request.urlopen(request_obj, timeout=8) as response:
                logo_bytes = response.read()
        except Exception:
            logo_bytes = None
    else:
        static_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'static'))
        logo_path = os.path.join(static_root, logo_url.lstrip('/'))
        if not os.path.exists(logo_path):
            logo_path = None

    if logo_bytes is None and logo_path is None:
        try:
            request_obj = urllib.request.Request(
                default_logo_url,
                headers={
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
                                  '(KHTML, like Gecko) Chrome/115.0 Safari/537.36'
                },
            )
            with urllib.request.urlopen(request_obj, timeout=8) as response:
                logo_bytes = response.read()
        except Exception:
            logo_bytes = None

    if logo_bytes is not None or logo_path is not None:
        section = doc.sections[0]
        section.different_first_page_header_footer = True
        header = section.first_page_header
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = header_paragraph.add_run()
        if logo_bytes is not None:
            run.add_picture(io.BytesIO(logo_bytes), width=Inches(2.5))
        else:
            run.add_picture(logo_path, width=Inches(2.5))
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(title, level=0)
    if company:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Компания"
        hdr[1].text = company.company_name or ""

        rows = [
            ("Краткое название", company.short_name),
            ("Юридическая форма", company.legal_form),
            ("ИНН/КПП", f"{company.inn or ''} / {company.kpp or ''}".strip()),
            ("ОГРН", company.ogrn),
            ("Адрес", company.legal_address),
            ("Телефон", company.phone),
            ("Email", company.email),
        ]
        for label, value in rows:
            if value:
                row = table.add_row().cells
                row[0].text = label
                row[1].text = str(value)
        doc.add_paragraph()

    if period:
        doc.add_paragraph(f"Период: {period}")
    if snapshot_date:
        doc.add_paragraph(f"Дата среза: {snapshot_date}")

    doc.add_paragraph()

    if now:
        footer = doc.sections[0].footer
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_paragraph.text = f"Формирование: {now.strftime('%Y-%m-%d %H:%M')}"
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


def add_key_value_section(doc, title, items):
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Показатель"
    hdr[1].text = "Значение"
    for label, value in items:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = format_value(value)
    doc.add_paragraph()


def add_table(doc, title, headers, rows):
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = format_value(value)
    doc.add_paragraph()


def add_signatures_section(doc, company, show_seal, show_signatures):
    if not show_signatures and not show_seal:
        return
    doc.add_heading("Подписи и подтверждения", level=1)
    if company and company.ceo_position and company.ceo:
        doc.add_paragraph(f"{company.ceo_position}: {company.ceo}")
    elif company and company.ceo:
        doc.add_paragraph(f"Руководитель: {company.ceo}")
    if company and company.chief_accountant_name:
        doc.add_paragraph(f"Главный бухгалтер: {company.chief_accountant_name}")
    if show_seal:
        doc.add_paragraph("Печать: ________________________________")
    if show_signatures:
        doc.add_paragraph("Подпись: ______________________________")
    doc.add_paragraph()


def save_docx_response(doc, filename):
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return send_file(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


def build_dashboard_docx(
    company,
    period,
    revenue,
    cogs,
    gross_profit,
    operating_expenses,
    net_profit,
    inventory_value,
    cash_amount,
    receivables,
    assets,
    liabilities,
    equity,
    roe,
    breakeven,
    strength_buffer,
    show_seal,
    show_signatures,
    now,
):
    doc = create_docx_document("Финансовый дашборд", company, period, now=now)
    add_key_value_section(
        doc,
        "Ключевые показатели",
        [
            ("Выручка", revenue),
            ("Себестоимость продаж", cogs),
            ("Валовая прибыль", gross_profit),
            ("Операционные расходы", operating_expenses),
            ("Чистая прибыль", net_profit),
            ("Запасы", inventory_value),
            ("Денежные средства", cash_amount),
            ("Дебиторская задолженность", receivables),
            ("Итого активы", assets),
            ("Итого обязательства", liabilities),
            ("Собственный капитал", equity),
            ("ROE", format_percent(roe)),
            ("Точка безубыточности", breakeven),
            ("Запас прочности", format_percent(strength_buffer)),
        ],
    )
    add_signatures_section(doc, company, show_seal, show_signatures)
    filename = f"{period}_dashboard_{now.strftime('%Y%m%d')}.docx"
    return save_docx_response(doc, filename)


def build_management_balance_docx(
    company,
    snapshot_date,
    inventory_value,
    cash_amount,
    receivables,
    payables,
    total_assets,
    total_liabilities,
    equity,
    check_equal,
    show_seal,
    show_signatures,
    now,
):
    doc = create_docx_document("Управленческий баланс", company, snapshot_date=snapshot_date, now=now)
    add_key_value_section(
        doc,
        "Структура баланса",
        [
            ("Запасы", inventory_value),
            ("Денежные средства", cash_amount),
            ("Дебиторская задолженность", receivables),
            ("Итого активы", total_assets),
            ("Кредиторская задолженность", payables),
            ("Собственный капитал", equity),
        ],
    )
    doc.add_paragraph(f"Баланс совпадает: {'Да' if check_equal else 'Нет'}")
    add_signatures_section(doc, company, show_seal, show_signatures)
    filename = f"management_balance_{snapshot_date.strftime('%Y%m%d')}.docx"
    return save_docx_response(doc, filename)


def build_plan_fact_analysis_docx(
    company,
    period,
    analysis,
    deviations,
    deviations_alpha,
    total_planned,
    total_actual,
    total_variance,
    cash_gap,
    show_seal,
    show_signatures,
    now,
):
    doc = create_docx_document("План-факт анализ", company, period, now=now)
    add_key_value_section(
        doc,
        "Итоги по периоду",
        [
            ("План total", total_planned),
            ("Факт total", total_actual),
            ("Отклонение", total_variance),
            ("Денежный разрыв", cash_gap),
        ],
    )
    add_table(
        doc,
        "Анализ по категориям",
        ["Категория", "Тип", "План", "Факт", "Отклонение", "Сглажено"],
        [
            [
                item["category"],
                item["type"],
                item["planned"],
                item["actual"],
                item["variance"],
                item["flexed"],
            ]
            for item in analysis
        ],
    )
    if deviations_alpha:
        doc.add_heading("Крупные отклонения", level=1)
        for item in deviations_alpha:
            doc.add_paragraph(
                f"{item['category']}: {format_percent(item['variance_pct'])}"
            )
        doc.add_paragraph()
    if deviations:
        doc.add_heading("Комментарии к отклонениям", level=1)
        for item in deviations:
            doc.add_paragraph(f"{item.category}: {getattr(item, 'description', '')}")
    add_signatures_section(doc, company, show_seal, show_signatures)
    filename = f"{period}_plan_fact_{now.strftime('%Y%m%d')}.docx"
    return save_docx_response(doc, filename)


def build_cash_flow_docx(
    company,
    period,
    active_rows,
    total_incoming,
    total_outgoing_paid,
    total_outgoing_scheduled,
    closing_balance,
    show_seal,
    show_signatures,
    now,
):
    doc = create_docx_document("Кассовый отчет", company, period, now=now)
    add_key_value_section(
        doc,
        "Итоги за период",
        [
            ("Итого поступлений", total_incoming),
            ("Итого оплаченных расходов", total_outgoing_paid),
            ("Итого запланированных расходов", total_outgoing_scheduled),
            ("Закрывающий баланс", closing_balance),
        ],
    )
    if active_rows:
        add_table(
            doc,
            "Движение денежных средств по дням",
            [
                "Дата",
                "Поступления",
                "Оплачено",
                "Запланировано",
                "Чистый поток",
                "Проектный разрыв",
                "Баланс",
            ],
            [
                [
                    row["date"].strftime("%Y-%m-%d"),
                    row["incoming"],
                    row["outgoing_paid"],
                    row["outgoing_scheduled"],
                    row["net_actual"],
                    row["projected_gap"],
                    row["running_balance"],
                ]
                for row in active_rows
            ],
        )
    else:
        doc.add_paragraph("Нет операций за выбранный период.")
    add_signatures_section(doc, company, show_seal, show_signatures)
    filename = f"{period}_cash_flow_{now.strftime('%Y%m%d')}.docx"
    return save_docx_response(doc, filename)


def build_bdds_forecast_docx(
    company,
    rows,
    ending_balance,
    critical_dates,
    show_seal,
    show_signatures,
    now,
):
    doc = create_docx_document("Прогноз денежных потоков на 30 дней", company, now=now)
    add_key_value_section(
        doc,
        "Ключевые параметры",
        [
            ("Конечный баланс", ending_balance),
            ("Критические даты", ", ".join(d.strftime("%Y-%m-%d") for d in critical_dates) if critical_dates else "Нет"),
        ],
    )
    if rows:
        add_table(
            doc,
            "Прогноз движения денежных средств",
            ["Дата", "Поступления", "Выплаты", "Чистый поток", "Баланс"],
            [
                [
                    row["date"].strftime("%Y-%m-%d"),
                    row["incoming"],
                    row["outgoing"],
                    row["net"],
                    row["balance"],
                ]
                for row in rows
            ],
        )
    else:
        doc.add_paragraph("Нет данных для прогноза за 30 дней.")
    add_signatures_section(doc, company, show_seal, show_signatures)
    filename = f"bdds_forecast_{now.strftime('%Y%m%d')}.docx"
    return save_docx_response(doc, filename)


def build_settlements_docx(
    company,
    period,
    customer_rows,
    supplier_rows,
    total_customer_billed,
    total_customer_paid,
    total_supplier_accrued,
    total_supplier_paid,
    top_customer_income,
    top_supplier_payables,
    show_seal,
    show_signatures,
    now,
):
    doc = create_docx_document("Взаиморасчеты", company, period, now=now)
    add_key_value_section(
        doc,
        "Итоги по расчетам",
        [
            ("Всего начислено клиентам", total_customer_billed),
            ("Всего оплачено клиентами", total_customer_paid),
            ("Всего начислено поставщикам", total_supplier_accrued),
            ("Всего оплачено поставщикам", total_supplier_paid),
        ],
    )
    if customer_rows:
        add_table(
            doc,
            "Расчеты по клиентам",
            ["Контрагент", "Тип", "Заказы", "Начислено", "Оплачено", "Долг"],
            [
                [
                    row["name"],
                    row["type"],
                    row["orders_count"],
                    row["billed"],
                    row["paid"],
                    row["debt"],
                ]
                for row in customer_rows
            ],
        )
    if supplier_rows:
        add_table(
            doc,
            "Расчеты по поставщикам",
            ["Контрагент", "Заказы", "Начислено", "Оплачено", "К выплате"],
            [
                [
                    row["name"],
                    row["orders_count"],
                    row["accrued"],
                    row["paid"],
                    row["payable"],
                ]
                for row in supplier_rows
            ],
        )
    if top_customer_income:
        add_table(
            doc,
            "Топ клиентов по выручке",
            ["Контрагент", "Выручка"],
            [[name, amount] for name, amount in top_customer_income],
        )
    if top_supplier_payables:
        add_table(
            doc,
            "Топ поставщиков по задолженности",
            ["Контрагент", "Задолженность"],
            [[name, amount] for name, amount in top_supplier_payables],
        )
    add_signatures_section(doc, company, show_seal, show_signatures)
    filename = f"{period}_settlements_{now.strftime('%Y%m%d')}.docx"
    return save_docx_response(doc, filename)


def build_budget_docx(
    company,
    period,
    income_items,
    expense_items,
    total_income_plan,
    total_expense_plan,
    planned_balance,
    show_seal,
    show_signatures,
    now,
):
    doc = create_docx_document("Бюджет доходов и расходов", company, period, now=now)
    add_key_value_section(
        doc,
        "Итоги бюджета",
        [
            ("План доходов", total_income_plan),
            ("План расходов", total_expense_plan),
            ("Плановый остаток", planned_balance),
        ],
    )
    if income_items:
        add_table(
            doc,
            "Доходы",
            ["Категория", "План"],
            [[item.category, item.planned_amount] for item in income_items],
        )
    if expense_items:
        add_table(
            doc,
            "Расходы",
            ["Категория", "План"],
            [[item.category, item.planned_amount] for item in expense_items],
        )
    add_signatures_section(doc, company, show_seal, show_signatures)
    filename = f"{period}_budget_{now.strftime('%Y%m%d')}.docx"
    return save_docx_response(doc, filename)


def build_bdr_docx(
    company,
    period,
    pivot,
    total_revenue,
    total_cogs,
    total_profit,
    show_seal,
    show_signatures,
    now,
):
    doc = create_docx_document("Анализ рентабельности (БДР)", company, period, now=now)
    add_key_value_section(
        doc,
        "Итоги сегментов",
        [
            ("Выручка", total_revenue),
            ("Себестоимость", total_cogs),
            ("Прибыль", total_profit),
        ],
    )
    add_table(
        doc,
        "Показатели по сегментам",
        [
            "Сегмент",
            "Группа",
            "Выручка",
            "Себестоимость",
            "Валовая прибыль",
            "Валовая маржа",
            "Наценка",
            "Количество",
        ],
        [
            [
                item["segment"],
                item["group"],
                item["revenue"],
                item["cogs"],
                item["gross_profit"],
                format_percent(item["gross_margin_pct"]),
                format_percent(item["markup_pct"]),
                item["quantity"],
            ]
            for item in pivot
        ],
    )
    add_signatures_section(doc, company, show_seal, show_signatures)
    filename = f"{period}_bdr_{now.strftime('%Y%m%d')}.docx"
    return save_docx_response(doc, filename)


def finance_required(view):
    @wraps(view)
    def wrapper(*args, **kwargs):
        if not current_user.is_authenticated or not (
            current_user.is_finance or current_user.is_admin
        ):
            flash(
                "Финансовый модуль доступен бухгалтерии и администраторам.",
                "danger",
            )
            return redirect(url_for("index"))
        return view(*args, **kwargs)

    return wrapper


def get_period_bounds(period):
    try:
        start_date = datetime.strptime(period, "%Y-%m")
    except ValueError:
        start_date = datetime.now().replace(
            day=1, hour=0, minute=0, second=0, microsecond=0
        )
    end_date = (start_date + timedelta(days=32)).replace(day=1)
    return start_date, end_date


def calculate_income(start_date, end_date):
    value = (
        db.session.query(func.sum(Payment.amount))
        .filter(
            Payment.payment_date >= start_date,
            Payment.payment_date < end_date,
            Payment.status == "completed",
        )
        .scalar()
    )
    return float(value or 0)


def calculate_actual_expenses(start_date, end_date):
    value = (
        db.session.query(func.sum(PurchaseOrderItem.quantity * PurchaseOrderItem.unit_cost))
        .join(PurchaseOrder)
        .filter(
            PurchaseOrder.order_date >= start_date,
            PurchaseOrder.order_date < end_date,
        )
        .scalar()
    )
    return float(value or 0)


def calculate_cogs_fifo(start_date, end_date, strategy="fifo"):
    cogs = 0.0
    batch_cache = {}

    sales_items = (
        SalesOrderItem.query.join(SalesOrder)
        .filter(
            SalesOrder.order_date >= start_date,
            SalesOrder.order_date < end_date,
            SalesOrder.status == "completed",
        )
        .all()
    )

    for item in sales_items:
        planned = float(item.quantity * item.cost_price)
        if item.cost_price and item.cost_price > 0:
            cogs += planned
            continue

        product_id = item.product_id
        if product_id not in batch_cache:
            batch_cache[product_id] = (
                InventoryBatch.query.filter_by(product_id=product_id)
                .order_by(InventoryBatch.received_date)
                .all()
            )

        remaining = item.quantity
        for batch in batch_cache[product_id]:
            available = batch.available_quantity()
            if available <= 0:
                continue
            take = min(remaining, available)
            batch_unit_cost = batch.unit_cost + (batch.transport_cost / batch.quantity if batch.quantity else 0)
            cogs += take * batch_unit_cost
            remaining -= take
            if remaining <= 0:
                break

        if remaining > 0:
            cogs += remaining * (item.cost_price or 0)

    return cogs


def calculate_period_customer_income(start_date, end_date):
    customer_income = defaultdict(float)
    payments = (
        Payment.query.join(SalesOrder)
        .filter(
            Payment.payment_date >= start_date,
            Payment.payment_date < end_date,
            Payment.status == "completed",
        )
        .all()
    )
    for payment in payments:
        customer_income[payment.sales_order.customer.name] += payment.amount
    return customer_income


def calculate_period_supplier_payables(start_date, end_date):
    supplier_payables = defaultdict(float)
    orders = PurchaseOrder.query.filter(
        PurchaseOrder.order_date >= start_date,
        PurchaseOrder.order_date < end_date,
    ).all()
    for order in orders:
        supplier_payables[order.supplier.name] += order.total_amount
    return supplier_payables


def get_indirect_expense_map(period):
    data = defaultdict(float)
    for item in IndirectExpense.query.filter_by(period=period).all():
        data[item.category] += item.amount
    return data


@finance_bp.route("/")
@login_required
@finance_required
def index():
    current_period = datetime.now().strftime("%Y-%m")
    start_date, end_date = get_period_bounds(current_period)
    total_income = calculate_income(start_date, end_date)
    total_expenses = calculate_actual_expenses(start_date, end_date)
    current_liquidity = total_income - total_expenses

    period_rows = []
    for period in ["2026-01", "2026-02", "2026-03"]:
        period_start, period_end = get_period_bounds(period)
        income = calculate_income(period_start, period_end)
        expenses = calculate_actual_expenses(period_start, period_end)
        period_rows.append(
            {
                "period": period,
                "income": income,
                "expenses": expenses,
                "balance": income - expenses,
            }
        )

    unpaid_purchase_orders = PurchaseOrder.query.filter_by(is_paid=False).count()
    unpaid_purchase_amount = sum(
        order.total_amount for order in PurchaseOrder.query.filter_by(is_paid=False).all()
    )

    return render_template(
        "finance/index.html",
        payments_count=Payment.query.count(),
        purchase_orders_count=PurchaseOrder.query.count(),
        sales_orders_count=SalesOrder.query.count(),
        total_income=total_income,
        total_expenses=total_expenses,
        current_liquidity=current_liquidity,
        current_period=current_period,
        period_rows=period_rows,
        unpaid_purchase_orders=unpaid_purchase_orders,
        unpaid_purchase_amount=unpaid_purchase_amount,
    )


@finance_bp.route("/incoming-payments")
@login_required
@finance_required
def incoming_payments():
    page = request.args.get("page", 1, type=int)
    date_from = request.args.get("date_from", "").strip()
    date_to = request.args.get("date_to", "").strip()
    q = request.args.get("q", "").strip()

    query = Payment.query.join(SalesOrder).join(Customer)
    if date_from:
        try:
            query = query.filter(Payment.payment_date >= datetime.strptime(date_from, "%Y-%m-%d"))
        except ValueError:
            flash("Некорректная дата начала периода.", "warning")
    if date_to:
        try:
            query = query.filter(
                Payment.payment_date < datetime.strptime(date_to, "%Y-%m-%d") + timedelta(days=1)
            )
        except ValueError:
            flash("Некорректная дата окончания периода.", "warning")
    if q:
        query = query.filter(
            or_(
                SalesOrder.order_number.ilike(f"%{q}%"),
                Payment.fiscal_receipt_number.ilike(f"%{q}%"),
                Customer.name.ilike(f"%{q}%"),
            )
        )

    payments = query.order_by(Payment.payment_date.desc()).paginate(page=page, per_page=20)
    filtered_payments = query.all()
    return render_template(
        "finance/payments/incoming.html",
        payments=payments,
        date_from=date_from,
        date_to=date_to,
        q=q,
        total_amount=sum(payment.amount for payment in filtered_payments),
        completed_count=sum(1 for payment in filtered_payments if payment.status == "completed"),
    )


@finance_bp.route("/outgoing-payments")
@login_required
@finance_required
def outgoing_payments():
    page = request.args.get("page", 1, type=int)
    q = request.args.get("q", "").strip()
    status = request.args.get("status", "").strip()

    query = PurchaseOrder.query.join(Supplier)
    if q:
        query = query.filter(
            or_(
                PurchaseOrder.order_number.ilike(f"%{q}%"),
                Supplier.name.ilike(f"%{q}%"),
            )
        )
    if status == "paid":
        query = query.filter(PurchaseOrder.is_paid.is_(True))
    elif status == "unpaid":
        query = query.filter(PurchaseOrder.is_paid.is_(False))

    purchase_orders = query.order_by(PurchaseOrder.order_date.desc()).paginate(page=page, per_page=20)
    filtered = query.all()
    return render_template(
        "finance/payments/outgoing.html",
        purchase_orders=purchase_orders,
        q=q,
        status=status,
        pending_amount=sum(order.total_amount for order in filtered if not order.is_paid),
        paid_amount=sum(order.total_amount for order in filtered if order.is_paid),
        pending_count=sum(1 for order in filtered if not order.is_paid),
    )


@finance_bp.route("/outgoing-payments/<int:po_id>/mark-paid", methods=["POST"])
@login_required
@finance_required
def mark_purchase_order_paid(po_id):
    purchase_order = PurchaseOrder.query.get_or_404(po_id)
    purchase_order.is_paid = True
    purchase_order.status = "completed"
    db.session.commit()
    flash(
        f"Заказ поставщику {purchase_order.order_number} отмечен как оплаченный.",
        "success",
    )
    return redirect(url_for("finance.outgoing_payments"))


@finance_bp.route("/budget")
@login_required
@finance_required
def budget():
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    items = (
        BudgetItem.query.filter_by(period=period)
        .order_by(BudgetItem.item_type, BudgetItem.category)
        .all()
    )
    income_items = [item for item in items if item.item_type == "income"]
    expense_items = [item for item in items if item.item_type == "expense"]
    total_income_plan = sum(item.planned_amount for item in income_items)
    total_expense_plan = sum(item.planned_amount for item in expense_items)

    return render_template(
        "finance/budget/index.html",
        period=period,
        income_items=income_items,
        expense_items=expense_items,
        current_period=datetime.now().strftime("%Y-%m"),
        total_income_plan=total_income_plan,
        total_expense_plan=total_expense_plan,
        planned_balance=total_income_plan - total_expense_plan,
    )


@finance_bp.route("/budget/add", methods=["GET", "POST"])
@login_required
@finance_required
def add_budget_item():
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))

    if request.method == "POST":
        category = request.form.get("category", "").strip()
        if not category:
            flash("Категория обязательна.", "danger")
            return redirect(url_for("finance.add_budget_item", period=period))

        budget_item = BudgetItem(
            period=request.form.get("period", period),
            item_type=request.form.get("item_type", "expense"),
            category=category,
            planned_amount=float(request.form.get("planned_amount", 0) or 0),
        )
        db.session.add(budget_item)
        db.session.commit()

        flash("Статья бюджета добавлена.", "success")
        return redirect(url_for("finance.budget", period=budget_item.period))

    return render_template("finance/budget/add.html", period=period)


@finance_bp.route("/plan-fact-analysis", methods=["GET", "POST"])
@login_required
@finance_required
def plan_fact_analysis():
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    start_date, end_date = get_period_bounds(period)
    budget_items = BudgetItem.query.filter_by(period=period).all()

    if request.method == "POST":
        item_name = request.form.get("item_name", "").strip()
        reason = request.form.get("reason", "").strip()
        deviation_val = float(request.form.get("deviation", 0) or 0)
        if item_name and reason:
            deviation = PlanFactDeviation(
                period=period,
                item_name=item_name,
                planned_value=float(request.form.get("planned", 0) or 0),
                actual_value=float(request.form.get("actual", 0) or 0),
                deviation=deviation_val,
                deviation_pct=float(request.form.get("deviation_pct", 0) or 0),
                reason=reason,
                entered_by=current_user.id,
            )
            db.session.add(deviation)
            db.session.commit()
            flash("Причина отклонения сохранена", "success")
        else:
            flash("Заполните имя статьи и причину отклонения", "warning")
        return redirect(url_for("finance.plan_fact_analysis", period=period))

    planned_by_category = defaultdict(float)
    for item in budget_items:
        planned_by_category[(item.item_type, item.category)] += item.planned_amount

    actual_income = calculate_income(start_date, end_date)
    actual_expenses = calculate_actual_expenses(start_date, end_date)
    indirect_map = get_indirect_expense_map(period)

    flex_factor = (
        actual_income / planned_by_category.get(("income", "Продажи"), actual_income or 1)
    ) if planned_by_category.get(("income", "Продажи"), 0) else 1

    analysis = [
        {
            "category": "Продажи",
            "type": "income",
            "planned": planned_by_category.get(("income", "Продажи"), 0),
            "actual": actual_income * 0.75,
            "variance": actual_income * 0.75 - planned_by_category.get(("income", "Продажи"), 0),
            "flexed": planned_by_category.get(("income", "Продажи"), 0) * flex_factor,
        },
        {
            "category": "Корпоративные проекты",
            "type": "income",
            "planned": planned_by_category.get(("income", "Корпоративные проекты"), 0),
            "actual": actual_income * 0.15,
            "variance": actual_income * 0.15 - planned_by_category.get(("income", "Корпоративные проекты"), 0),
            "flexed": planned_by_category.get(("income", "Корпоративные проекты"), 0) * flex_factor,
        },
        {
            "category": "Сборка и доставка",
            "type": "income",
            "planned": planned_by_category.get(("income", "Сборка и доставка"), 0),
            "actual": actual_income * 0.10,
            "variance": actual_income * 0.10 - planned_by_category.get(("income", "Сборка и доставка"), 0),
            "flexed": planned_by_category.get(("income", "Сборка и доставка"), 0) * flex_factor,
        },
        {
            "category": "Закупки",
            "type": "expense",
            "planned": planned_by_category.get(("expense", "Закупки"), 0),
            "actual": actual_expenses,
            "variance": actual_expenses - planned_by_category.get(("expense", "Закупки"), 0),
            "flexed": planned_by_category.get(("expense", "Закупки"), 0) * flex_factor,
        },
    ]

    for (item_type, category), planned_value in sorted(planned_by_category.items()):
        if category in {"Продажи", "Корпоративные проекты", "Сборка и доставка", "Закупки"}:
            continue
        if item_type == "income":
            if category == "Корпоративные проекты":
                actual_value = actual_income * 0.15
            elif category == "Сборка и доставка":
                actual_value = actual_income * 0.10
            else:
                actual_value = actual_income * 0.05
        else:
            actual_value = float(indirect_map.get(category, 0))
        flexed_value = planned_value * flex_factor
        variance = actual_value - planned_value
        analysis.append(
            {
                "category": category,
                "type": item_type,
                "planned": planned_value,
                "actual": actual_value,
                "variance": variance,
                "flexed": flexed_value,
            }
        )

    deviations = PlanFactDeviation.query.filter_by(period=period).order_by(PlanFactDeviation.created_at.desc()).all()
    total_planned = sum(item["planned"] for item in analysis)
    total_actual = sum(item["actual"] for item in analysis)

    deviations_alpha = []
    for item in analysis:
        planned = item.get("planned", 0) or 0
        variance_pct = (item.get("variance", 0) / planned * 100) if planned else 0
        if abs(variance_pct) >= 20:
            deviations_alpha.append({"category": item["category"], "variance_pct": variance_pct})

    return render_template(
        "finance/plan_fact.html",
        period=period,
        start_date=start_date,
        end_date=end_date,
        analysis=analysis,
        deviations=deviations,
        deviations_alpha=deviations_alpha,
        total_planned=total_planned,
        total_actual=total_actual,
        total_variance=total_actual - total_planned,
        cash_gap=actual_income - actual_expenses,
    )


@finance_bp.route("/profitability-report")
@login_required
@finance_required
def profitability_report():
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    start_date, end_date = get_period_bounds(period)

    revenue = calculate_income(start_date, end_date)
    cogs = (
        db.session.query(func.sum(SalesOrderItem.quantity * SalesOrderItem.cost_price))
        .join(SalesOrder)
        .filter(
            SalesOrder.order_date >= start_date,
            SalesOrder.order_date < end_date,
            SalesOrder.status == "completed",
        )
        .scalar()
        or 0
    )
    cogs = float(cogs)
    gross_profit = revenue - cogs
    operating_expenses = calculate_actual_expenses(start_date, end_date)
    usn_tax = revenue * 0.06
    operating_profit = gross_profit - operating_expenses
    net_profit = operating_profit - usn_tax

    return render_template(
        "finance/profitability.html",
        period=period,
        start_date=start_date,
        end_date=end_date,
        revenue=revenue,
        cogs=cogs,
        gross_profit=gross_profit,
        operating_expenses=operating_expenses,
        operating_profit=operating_profit,
        usn_tax=usn_tax,
        net_profit=net_profit,
        gross_margin_pct=(gross_profit / revenue * 100) if revenue else 0,
        net_margin_pct=(net_profit / revenue * 100) if revenue else 0,
        cogs_note="Для MVP себестоимость фиксируется в момент продажи в поле cost_price.",
    )


@finance_bp.route("/indirect-expenses", methods=["GET", "POST"])
@login_required
@finance_required
def indirect_expenses():
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    if request.method == "POST":
        category = request.form.get("category", "").strip()
        amount = float(request.form.get("amount", 0) or 0)
        description = request.form.get("description", "").strip()
        if not category or amount <= 0:
            flash("Категория и сумма обязательны.", "warning")
            return redirect(url_for("finance.indirect_expenses", period=period))
        exp = IndirectExpense(period=period, category=category, amount=amount, description=description)
        db.session.add(exp)
        db.session.commit()
        flash("Косвенная статья добавлена.", "success")
        return redirect(url_for("finance.indirect_expenses", period=period))

    items = IndirectExpense.query.filter_by(period=period).order_by(IndirectExpense.category.asc()).all()
    return render_template("finance/indirect_expenses.html", period=period, items=items)


@finance_bp.route("/indirect-expenses/<int:expense_id>/delete", methods=["POST"])
@login_required
@finance_required
def delete_indirect_expense(expense_id):
    expense = IndirectExpense.query.get_or_404(expense_id)
    period = expense.period
    db.session.delete(expense)
    db.session.commit()
    flash("Косвенная статья удалена.", "success")
    return redirect(url_for("finance.indirect_expenses", period=period))


@finance_bp.route("/cash-calendar", methods=["GET", "POST"])
@login_required
@finance_required
def cash_calendar():
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    start_date, end_date = get_period_bounds(period)

    if request.method == "POST":
        date_str = request.form.get("date")
        amount = float(request.form.get("amount", 0) or 0)
        direction = request.form.get("direction")
        cash_type = request.form.get("cash_type")
        counterparty_id = request.form.get("counterparty_id") or None
        supplier_id = request.form.get("supplier_id") or None
        probability = float(request.form.get("probability", 1.0) or 1.0)
        comment = request.form.get("comment", "").strip()

        try:
            date = datetime.strptime(date_str, "%Y-%m-%d")
        except (ValueError, TypeError):
            flash("Неправильная дата", "warning")
            return redirect(url_for("finance.cash_calendar", period=period))

        item = CashCalendarItem(
            date=date,
            amount=amount,
            direction=direction,
            cash_type=cash_type,
            counterparty_id=counterparty_id,
            supplier_id=supplier_id,
            status="planned",
            probability=probability,
            comment=comment,
        )
        db.session.add(item)
        db.session.commit()
        flash("Элемент календаря сохранён.", "success")
        return redirect(url_for("finance.cash_calendar", period=period))

    items = CashCalendarItem.query.filter(CashCalendarItem.date >= start_date, CashCalendarItem.date < end_date).order_by(CashCalendarItem.date).all()
    return render_template("finance/cash_calendar.html", period=period, items=items)


@finance_bp.route("/cash-calendar/<int:item_id>/delete", methods=["POST"])
@login_required
@finance_required
def delete_cash_calendar_item(item_id):
    item = CashCalendarItem.query.get_or_404(item_id)
    period = item.date.strftime("%Y-%m")
    db.session.delete(item)
    db.session.commit()
    flash("Элемент календаря удалён.", "success")
    return redirect(url_for("finance.cash_calendar", period=period))


@finance_bp.route("/cash-flow")
@login_required
@finance_required
def cash_flow():
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    start_date, end_date = get_period_bounds(period)

    incoming_map = defaultdict(float)
    for payment in Payment.query.filter(
        Payment.payment_date >= start_date,
        Payment.payment_date < end_date,
        Payment.status == "completed",
    ).all():
        incoming_map[payment.payment_date.date()] += payment.amount

    outgoing_paid_map = defaultdict(float)
    outgoing_scheduled_map = defaultdict(float)
    for order in PurchaseOrder.query.filter(
        PurchaseOrder.order_date >= start_date,
        PurchaseOrder.order_date < end_date,
    ).all():
        if order.is_paid:
            outgoing_paid_map[order.order_date.date()] += order.total_amount
        else:
            outgoing_scheduled_map[order.order_date.date()] += order.total_amount

    rows = []
    running_balance = 0
    day = start_date
    while day < end_date:
        day_key = day.date()
        incoming = incoming_map.get(day_key, 0)
        outgoing_paid = outgoing_paid_map.get(day_key, 0)
        outgoing_scheduled = outgoing_scheduled_map.get(day_key, 0)
        running_balance += incoming - outgoing_paid
        rows.append(
            {
                "date": day_key,
                "incoming": incoming,
                "outgoing_paid": outgoing_paid,
                "outgoing_scheduled": outgoing_scheduled,
                "net_actual": incoming - outgoing_paid,
                "projected_gap": incoming - outgoing_paid - outgoing_scheduled,
                "running_balance": running_balance,
            }
        )
        day += timedelta(days=1)

    active_rows = [
        row
        for row in rows
        if row["incoming"] or row["outgoing_paid"] or row["outgoing_scheduled"]
    ]

    return render_template(
        "finance/cash_flow.html",
        period=period,
        rows=rows,
        active_rows=active_rows,
        total_incoming=sum(row["incoming"] for row in rows),
        total_outgoing_paid=sum(row["outgoing_paid"] for row in rows),
        total_outgoing_scheduled=sum(row["outgoing_scheduled"] for row in rows),
        closing_balance=running_balance,
    )


@finance_bp.route("/settlements")
@login_required
@finance_required
def settlements():
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    start_date, end_date = get_period_bounds(period)

    customer_rows = []
    for customer in Customer.query.order_by(Customer.name.asc()).all():
        orders = [
            order
            for order in customer.sales_orders
            if start_date <= order.order_date < end_date
        ]
        if not orders:
            continue
        billed = sum(order.total_amount for order in orders)
        paid = sum(
            payment.amount
            for order in orders
            for payment in order.payments
            if start_date <= payment.payment_date < end_date
            and payment.status == "completed"
        )
        customer_rows.append(
            {
                "name": customer.name,
                "type": customer.type,
                "orders_count": len(orders),
                "billed": billed,
                "paid": paid,
                "debt": billed - paid,
            }
        )

    supplier_rows = []
    for supplier in Supplier.query.order_by(Supplier.name.asc()).all():
        orders = [
            order
            for order in supplier.purchase_orders
            if start_date <= order.order_date < end_date
        ]
        if not orders:
            continue
        accrued = sum(order.total_amount for order in orders)
        paid = sum(order.total_amount for order in orders if order.is_paid)
        supplier_rows.append(
            {
                "name": supplier.name,
                "orders_count": len(orders),
                "accrued": accrued,
                "paid": paid,
                "payable": accrued - paid,
            }
        )

    customer_rows.sort(key=lambda row: row["debt"], reverse=True)
    supplier_rows.sort(key=lambda row: row["payable"], reverse=True)

    return render_template(
        "finance/settlements.html",
        period=period,
        customer_rows=customer_rows,
        supplier_rows=supplier_rows,
        total_customer_billed=sum(row["billed"] for row in customer_rows),
        total_customer_paid=sum(row["paid"] for row in customer_rows),
        total_supplier_accrued=sum(row["accrued"] for row in supplier_rows),
        total_supplier_paid=sum(row["paid"] for row in supplier_rows),
        top_customer_income=sorted(
            calculate_period_customer_income(start_date, end_date).items(),
            key=lambda item: item[1],
            reverse=True,
        )[:5],
        top_supplier_payables=sorted(
            calculate_period_supplier_payables(start_date, end_date).items(),
            key=lambda item: item[1],
            reverse=True,
        )[:5],
    )


def calculate_bdr_pivot(start_date, end_date):
    data = []
    line_items = (
        SalesOrderItem.query.join(SalesOrder)
        .filter(
            SalesOrder.order_date >= start_date,
            SalesOrder.order_date < end_date,
            SalesOrder.status == "completed",
        )
        .all()
    )

    pivot = defaultdict(lambda: {"revenue": 0.0, "cogs": 0.0, "quantity": 0})
    for item in line_items:
        segment = item.sales_order.segment or "retail"
        group = item.product_group or "Общие"
        rev = item.quantity * item.unit_price
        cogs = item.quantity * item.cost_price
        pivot[(segment, group)]["revenue"] += rev
        pivot[(segment, group)]["cogs"] += cogs
        pivot[(segment, group)]["quantity"] += item.quantity

    for (segment, group), values in pivot.items():
        gp = values["revenue"] - values["cogs"]
        data.append(
            {
                "segment": segment,
                "group": group,
                "revenue": values["revenue"],
                "cogs": values["cogs"],
                "gross_profit": gp,
                "gross_margin_pct": (gp / values["revenue"] * 100) if values["revenue"] else 0,
                "markup_pct": (gp / values["cogs"] * 100) if values["cogs"] else 0,
                "quantity": values["quantity"],
            }
        )
    return data


@finance_bp.route("/bdr")
@login_required
@finance_required
def bdr_report():
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    start_date, end_date = get_period_bounds(period)
    pivot = calculate_bdr_pivot(start_date, end_date)
    total_revenue = sum(x["revenue"] for x in pivot)
    total_cogs = calculate_cogs_fifo(start_date, end_date)
    total_profit = total_revenue - total_cogs

    # Сводная строка по сегментам и группам
    return render_template(
        "finance/bdr.html",
        period=period,
        pivot=pivot,
        total_revenue=total_revenue,
        total_cogs=total_cogs,
        total_profit=total_profit,
    )


def _build_cash_calendar_rows(start_date, end_date):
    incoming = defaultdict(float)
    outgoing = defaultdict(float)

    # Фактические платежи, уже произведенные
    for payment in Payment.query.filter(
        Payment.payment_date >= start_date,
        Payment.payment_date < end_date,
    ).all():
        if payment.status == "completed":
            incoming[payment.payment_date.date()] += payment.amount

    # Фактические расходы (оплаченные закупки)
    for order in PurchaseOrder.query.filter(
        PurchaseOrder.order_date >= start_date,
        PurchaseOrder.order_date < end_date,
    ).all():
        if order.is_paid:
            outgoing[order.order_date.date()] += order.total_amount

    # Плановые платежи для неоплаченных закупок (30 дней от заказа)
    for order in PurchaseOrder.query.filter(PurchaseOrder.order_date < end_date).all():
        if not order.is_paid:
            pay_date = (order.order_date + timedelta(days=30)).date()
            if pay_date < end_date.date():
                outgoing[pay_date] += order.total_amount * 0.95

    # Ожидаемые поступления от дебиторки (14 дней от заказа)
    for so in SalesOrder.query.filter(SalesOrder.order_date < end_date).all():
        accrued = so.total_amount or 0
        paid = sum(p.amount for p in so.payments if p.status == "completed" and p.payment_date < end_date)
        due = max(0, accrued - paid)
        if due > 0:
            collect_date = (so.order_date + timedelta(days=14)).date()
            if collect_date < end_date.date():
                incoming[collect_date] += due * 0.8

    # Позиции платежного календаря
    for item in CashCalendarItem.query.filter(
        CashCalendarItem.date >= start_date,
        CashCalendarItem.date < end_date,
    ).all():
        d = item.date.date()
        if item.direction == "incoming":
            incoming[d] += item.amount * item.probability
        else:
            outgoing[d] += item.amount * item.probability

    # Начальный остаток = сумма полученных платежей до начала периода
    running_balance = (
        db.session.query(func.sum(Payment.amount))
        .filter(Payment.status == "completed", Payment.payment_date < start_date)
        .scalar()
        or 0
    )

    rows = []
    day = start_date
    while day < end_date:
        d = day.date()
        inc = incoming.get(d, 0)
        out = outgoing.get(d, 0)
        running_balance += inc - out
        rows.append(
            {
                "date": d,
                "incoming": inc,
                "outgoing": out,
                "net": inc - out,
                "balance": running_balance,
                "is_gap": running_balance < 0,
            }
        )
        day += timedelta(days=1)

    return rows


@finance_bp.route("/bdds/calendar")
@login_required
@finance_required
def bdds_calendar():
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    start_date, end_date = get_period_bounds(period)
    rows = _build_cash_calendar_rows(start_date, end_date)
    closing_balance = rows[-1]["balance"] if rows else 0

    return render_template(
        "finance/bdds_calendar.html",
        period=period,
        rows=rows,
        closing_balance=closing_balance,
    )


@finance_bp.route("/bdds/forecast")
@login_required
@finance_required
def bdds_forecast():
    date_str = request.args.get("start_date", datetime.now().strftime("%Y-%m-%d"))
    try:
        start_date = datetime.strptime(date_str, "%Y-%m-%d")
    except ValueError:
        start_date = datetime.now()
    start_date = datetime.combine(start_date.date(), datetime.min.time())
    end_date = start_date + timedelta(days=30)
    rows = _build_cash_calendar_rows(start_date, end_date)

    if rows:
        critical_dates = [r["date"] for r in rows if r["balance"] < 0]
    else:
        critical_dates = []

    return render_template(
        "finance/bdds_forecast.html",
        rows=rows,
        start_date=start_date.date().isoformat(),
        end_date=end_date.date().isoformat(),
        ending_balance=rows[-1]["balance"] if rows else 0,
        critical_dates=critical_dates,
    )


@finance_bp.route("/help")
@login_required
@finance_required
def finance_help():
    return render_template("finance/help.html")


@finance_bp.route("/management-balance")
@login_required
@finance_required
def management_balance():
    date_str = request.args.get("date", datetime.now().strftime("%Y-%m-%d"))
    try:
        snapshot_date = datetime.strptime(date_str, "%Y-%m-%d")
    except ValueError:
        snapshot_date = datetime.now()

    inventory_value = 0
    for batch in InventoryBatch.query.all():
        inventory_value += batch.available_quantity() * batch.unit_cost

    cash_amount = (
        db.session.query(func.sum(Payment.amount))
        .filter(Payment.status == "completed", Payment.payment_date <= snapshot_date)
        .scalar()
        or 0
    )

    receivables = 0
    for order in SalesOrder.query.filter(SalesOrder.order_date <= snapshot_date).all():
        accrued = order.total_amount
        paid = sum(p.amount for p in order.payments if p.status == "completed" and p.payment_date <= snapshot_date)
        receivables += max(0, accrued - paid)

    payables = 0
    for order in PurchaseOrder.query.filter(PurchaseOrder.order_date <= snapshot_date).all():
        if not order.is_paid:
            payables += order.total_amount

    total_assets = inventory_value + cash_amount + receivables
    total_liabilities = payables
    equity = total_assets - total_liabilities

    return render_template(
        "finance/management_balance.html",
        snapshot_date=snapshot_date.date(),
        inventory_value=inventory_value,
        cash_amount=cash_amount,
        receivables=receivables,
        payables=payables,
        total_assets=total_assets,
        total_liabilities=total_liabilities,
        equity=equity,
        check_equal=(abs(total_assets - total_liabilities - equity) < 1e-2),
    )


@finance_bp.route("/dashboard")
@login_required
@finance_required
def dashboard():
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    start_date, end_date = get_period_bounds(period)

    revenue = calculate_income(start_date, end_date)
    cogs = (
        db.session.query(func.sum(SalesOrderItem.quantity * SalesOrderItem.cost_price))
        .join(SalesOrder)
        .filter(
            SalesOrder.order_date >= start_date,
            SalesOrder.order_date < end_date,
            SalesOrder.status == "completed",
        )
        .scalar()
        or 0
    )
    gross_profit = revenue - cogs
    operating_expenses = calculate_actual_expenses(start_date, end_date)
    net_profit = gross_profit - operating_expenses

    cash_flow_end = (
        db.session.query(func.sum(Payment.amount))
        .filter(Payment.status == "completed", Payment.payment_date < end_date)
        .scalar()
        or 0
    )

    # Составим управленческий набор статей активов аналогично management_balance
    inventory_value = sum(batch.available_quantity() * batch.unit_cost for batch in InventoryBatch.query.all())
    cash_amount = (
        db.session.query(func.sum(Payment.amount))
        .filter(Payment.status == "completed", Payment.payment_date < end_date)
        .scalar()
        or 0
    )

    receivables = 0
    for order in SalesOrder.query.filter(SalesOrder.order_date <= end_date).all():
        accrued = order.total_amount
        paid = sum(p.amount for p in order.payments if p.status == "completed" and p.payment_date <= end_date)
        receivables += max(0, accrued - paid)

    assets = total_assets = inventory_value + cash_amount + receivables
    liabilities = payables = sum(order.total_amount for order in PurchaseOrder.query.filter_by(is_paid=False).all())
    equity = total_assets - liabilities
    roe = (net_profit * 12 / equity * 100) if equity else 0

    fixed_costs = sum(exp.amount for exp in IndirectExpense.query.filter_by(period=period).all())
    breakeven = (fixed_costs / (gross_profit / revenue) if revenue and gross_profit > 0 else 0)
    strength_buffer = ((revenue - breakeven) / revenue * 100) if revenue else 0

    return render_template(
        "finance/dashboard.html",
        period=period,
        net_profit=net_profit,
        cash_balance=cash_flow_end,
        roe=roe,
        breakeven=breakeven,
        strength_buffer=strength_buffer,
        revenue=revenue,
        inventory_value=inventory_value,
        payables=payables,
        assets=assets,
        liabilities=liabilities,
        equity=equity,
    )


@finance_bp.route("/export-1c", methods=["POST"])
@login_required
@finance_required
def export_1c():
    period = request.form.get("period", datetime.now().strftime("%Y-%m"))
    start_date, end_date = get_period_bounds(period)

    sales_orders = SalesOrder.query.filter(
        SalesOrder.order_date >= start_date,
        SalesOrder.order_date < end_date,
    ).all()
    purchase_orders = PurchaseOrder.query.filter(
        PurchaseOrder.order_date >= start_date,
        PurchaseOrder.order_date < end_date,
    ).all()

    payload = {
        "export_date": datetime.now().isoformat(),
        "period": period,
        "sales_orders": [],
        "purchase_orders": [],
    }

    for order in sales_orders:
        payload["sales_orders"].append(
            {
                "order_number": order.order_number,
                "customer": order.customer.name,
                "order_date": order.order_date.isoformat(),
                "status": order.status,
                "total_amount": order.total_amount,
                "items": [
                    {
                        "sku": item.product.sku,
                        "name": item.product.name,
                        "quantity": item.quantity,
                        "unit_price": item.unit_price,
                        "cost_price": item.cost_price,
                    }
                    for item in order.items
                ],
            }
        )

    for order in purchase_orders:
        payload["purchase_orders"].append(
            {
                "order_number": order.order_number,
                "supplier": order.supplier.name,
                "order_date": order.order_date.isoformat(),
                "status": order.status,
                "total_amount": order.total_amount,
                "is_paid": order.is_paid,
                "items": [
                    {
                        "sku": item.product.sku,
                        "name": item.product.name,
                        "quantity": item.quantity,
                        "unit_cost": item.unit_cost,
                    }
                    for item in order.items
                ],
            }
        )

    content = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    return send_file(
        io.BytesIO(content),
        mimetype="application/json",
        as_attachment=True,
        download_name=f"mebelgrad_exchange_{period}.json",
    )


# ============================================================================
# PRINT ENDPOINTS / ПЕЧАТЬ ОТЧЁТОВ
# ============================================================================

@finance_bp.route("/dashboard/print")
@login_required
@finance_required
def dashboard_print():
    """Печать финансового дашборда"""
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    show_seal = request.args.get("show_seal", "1") in ["1", "true", "on"]
    show_signatures = request.args.get("show_signatures", "1") in ["1", "true", "on"]
    
    start_date, end_date = get_period_bounds(period)
    
    revenue = calculate_income(start_date, end_date)
    cogs = (
        db.session.query(func.sum(SalesOrderItem.quantity * SalesOrderItem.cost_price))
        .join(SalesOrder)
        .filter(
            SalesOrder.order_date >= start_date,
            SalesOrder.order_date < end_date,
            SalesOrder.status == "completed",
        )
        .scalar()
        or 0
    )
    gross_profit = revenue - cogs
    operating_expenses = calculate_actual_expenses(start_date, end_date)
    net_profit = gross_profit - operating_expenses
    
    inventory_value = sum(batch.available_quantity() * batch.unit_cost for batch in InventoryBatch.query.all())
    cash_amount = (
        db.session.query(func.sum(Payment.amount))
        .filter(Payment.status == "completed", Payment.payment_date < end_date)
        .scalar()
        or 0
    )
    
    receivables = 0
    for order in SalesOrder.query.filter(SalesOrder.order_date <= end_date).all():
        accrued = order.total_amount
        paid = sum(p.amount for p in order.payments if p.status == "completed" and p.payment_date <= end_date)
        receivables += max(0, accrued - paid)
    
    assets = total_assets = inventory_value + cash_amount + receivables
    liabilities = payables = sum(order.total_amount for order in PurchaseOrder.query.filter_by(is_paid=False).all())
    equity = total_assets - liabilities
    roe = (net_profit * 12 / equity * 100) if equity else 0
    
    fixed_costs = sum(exp.amount for exp in IndirectExpense.query.filter_by(period=period).all())
    breakeven = (fixed_costs / (gross_profit / revenue) if revenue and gross_profit > 0 else 0)
    strength_buffer = ((revenue - breakeven) / revenue * 100) if revenue else 0
    
    company = CompanyProfile.query.first()
    
    return build_dashboard_docx(
        company=company,
        period=period,
        revenue=revenue,
        cogs=cogs,
        gross_profit=gross_profit,
        operating_expenses=operating_expenses,
        net_profit=net_profit,
        inventory_value=inventory_value,
        cash_amount=cash_amount,
        receivables=receivables,
        assets=assets,
        liabilities=liabilities,
        equity=equity,
        roe=roe,
        breakeven=breakeven,
        strength_buffer=strength_buffer,
        show_seal=show_seal,
        show_signatures=show_signatures,
        now=datetime.now(),
    )


@finance_bp.route("/management-balance/print")
@login_required
@finance_required
def management_balance_print():
    """Печать баланса по управленческому учету"""
    date_str = request.args.get("date", datetime.now().strftime("%Y-%m-%d"))
    show_seal = request.args.get("show_seal", "1") in ["1", "true", "on"]
    show_signatures = request.args.get("show_signatures", "1") in ["1", "true", "on"]
    
    try:
        snapshot_date = datetime.strptime(date_str, "%Y-%m-%d")
    except ValueError:
        snapshot_date = datetime.now()
    
    inventory_value = 0
    for batch in InventoryBatch.query.all():
        inventory_value += batch.available_quantity() * batch.unit_cost
    
    cash_amount = (
        db.session.query(func.sum(Payment.amount))
        .filter(Payment.status == "completed", Payment.payment_date <= snapshot_date)
        .scalar()
        or 0
    )
    
    receivables = 0
    for order in SalesOrder.query.filter(SalesOrder.order_date <= snapshot_date).all():
        accrued = order.total_amount
        paid = sum(p.amount for p in order.payments if p.status == "completed" and p.payment_date <= snapshot_date)
        receivables += max(0, accrued - paid)
    
    payables = 0
    for order in PurchaseOrder.query.filter(PurchaseOrder.order_date <= snapshot_date).all():
        if not order.is_paid:
            payables += order.total_amount
    
    total_assets = inventory_value + cash_amount + receivables
    total_liabilities = payables
    equity = total_assets - total_liabilities
    
    company = CompanyProfile.query.first()
    
    return build_management_balance_docx(
        company=company,
        snapshot_date=snapshot_date.date(),
        inventory_value=inventory_value,
        cash_amount=cash_amount,
        receivables=receivables,
        payables=payables,
        total_assets=total_assets,
        total_liabilities=total_liabilities,
        equity=equity,
        check_equal=(abs(total_assets - total_liabilities - equity) < 1e-2),
        show_seal=show_seal,
        show_signatures=show_signatures,
        now=datetime.now(),
    )


@finance_bp.route("/plan-fact-analysis/print")
@login_required
@finance_required
def plan_fact_analysis_print():
    """Печать анализа плана и факта"""
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    show_seal = request.args.get("show_seal", "1") in ["1", "true", "on"]
    show_signatures = request.args.get("show_signatures", "1") in ["1", "true", "on"]
    
    start_date, end_date = get_period_bounds(period)
    budget_items = BudgetItem.query.filter_by(period=period).all()
    
    planned_by_category = defaultdict(float)
    for item in budget_items:
        planned_by_category[(item.item_type, item.category)] += item.planned_amount
    
    actual_income = calculate_income(start_date, end_date)
    actual_expenses = calculate_actual_expenses(start_date, end_date)
    indirect_map = get_indirect_expense_map(period)
    
    flex_factor = (
        actual_income / planned_by_category.get(("income", "Продажи"), actual_income or 1)
    ) if planned_by_category.get(("income", "Продажи"), 0) else 1
    
    analysis = [
        {
            "category": "Продажи",
            "type": "income",
            "planned": planned_by_category.get(("income", "Продажи"), 0),
            "actual": actual_income * 0.75,
            "variance": actual_income * 0.75 - planned_by_category.get(("income", "Продажи"), 0),
            "flexed": planned_by_category.get(("income", "Продажи"), 0) * flex_factor,
        },
        {
            "category": "Корпоративные проекты",
            "type": "income",
            "planned": planned_by_category.get(("income", "Корпоративные проекты"), 0),
            "actual": actual_income * 0.15,
            "variance": actual_income * 0.15 - planned_by_category.get(("income", "Корпоративные проекты"), 0),
            "flexed": planned_by_category.get(("income", "Корпоративные проекты"), 0) * flex_factor,
        },
        {
            "category": "Сборка и доставка",
            "type": "income",
            "planned": planned_by_category.get(("income", "Сборка и доставка"), 0),
            "actual": actual_income * 0.10,
            "variance": actual_income * 0.10 - planned_by_category.get(("income", "Сборка и доставка"), 0),
            "flexed": planned_by_category.get(("income", "Сборка и доставка"), 0) * flex_factor,
        },
        {
            "category": "Закупки",
            "type": "expense",
            "planned": planned_by_category.get(("expense", "Закупки"), 0),
            "actual": actual_expenses,
            "variance": actual_expenses - planned_by_category.get(("expense", "Закупки"), 0),
            "flexed": planned_by_category.get(("expense", "Закупки"), 0) * flex_factor,
        },
    ]
    
    for (item_type, category), planned_value in sorted(planned_by_category.items()):
        if category in {"Продажи", "Корпоративные проекты", "Сборка и доставка", "Закупки"}:
            continue
        if item_type == "income":
            if category == "Корпоративные проекты":
                actual_value = actual_income * 0.15
            elif category == "Сборка и доставка":
                actual_value = actual_income * 0.10
            else:
                actual_value = actual_income * 0.05
        else:
            actual_value = float(indirect_map.get(category, 0))
        flexed_value = planned_value * flex_factor
        variance = actual_value - planned_value
        analysis.append(
            {
                "category": category,
                "type": item_type,
                "planned": planned_value,
                "actual": actual_value,
                "variance": variance,
                "flexed": flexed_value,
            }
        )
    
    deviations = PlanFactDeviation.query.filter_by(period=period).order_by(PlanFactDeviation.created_at.desc()).all()
    total_planned = sum(item["planned"] for item in analysis)
    total_actual = sum(item["actual"] for item in analysis)
    
    deviations_alpha = []
    for item in analysis:
        planned = item.get("planned", 0) or 0
        variance_pct = (item.get("variance", 0) / planned * 100) if planned else 0
        if abs(variance_pct) >= 20:
            deviations_alpha.append({"category": item["category"], "variance_pct": variance_pct})
    
    company = CompanyProfile.query.first()
    
    return build_plan_fact_analysis_docx(
        company=company,
        period=period,
        analysis=analysis,
        deviations=deviations,
        deviations_alpha=deviations_alpha,
        total_planned=total_planned,
        total_actual=total_actual,
        total_variance=total_actual - total_planned,
        cash_gap=actual_income - actual_expenses,
        show_seal=show_seal,
        show_signatures=show_signatures,
        now=datetime.now(),
    )


@finance_bp.route("/cash-flow/print")
@login_required
@finance_required
def cash_flow_print():
    """Печать отчета о кассовых потоках"""
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    show_seal = request.args.get("show_seal", "1") in ["1", "true", "on"]
    show_signatures = request.args.get("show_signatures", "1") in ["1", "true", "on"]
    
    start_date, end_date = get_period_bounds(period)
    
    incoming_map = defaultdict(float)
    for payment in Payment.query.filter(
        Payment.payment_date >= start_date,
        Payment.payment_date < end_date,
        Payment.status == "completed",
    ).all():
        incoming_map[payment.payment_date.date()] += payment.amount
    
    outgoing_paid_map = defaultdict(float)
    outgoing_scheduled_map = defaultdict(float)
    for order in PurchaseOrder.query.filter(
        PurchaseOrder.order_date >= start_date,
        PurchaseOrder.order_date < end_date,
    ).all():
        if order.is_paid:
            outgoing_paid_map[order.order_date.date()] += order.total_amount
        else:
            outgoing_scheduled_map[order.order_date.date()] += order.total_amount
    
    rows = []
    running_balance = 0
    day = start_date
    while day < end_date:
        day_key = day.date()
        incoming = incoming_map.get(day_key, 0)
        outgoing_paid = outgoing_paid_map.get(day_key, 0)
        outgoing_scheduled = outgoing_scheduled_map.get(day_key, 0)
        running_balance += incoming - outgoing_paid
        rows.append(
            {
                "date": day_key,
                "incoming": incoming,
                "outgoing_paid": outgoing_paid,
                "outgoing_scheduled": outgoing_scheduled,
                "net_actual": incoming - outgoing_paid,
                "projected_gap": incoming - outgoing_paid - outgoing_scheduled,
                "running_balance": running_balance,
            }
        )
        day += timedelta(days=1)
    
    active_rows = [
        row
        for row in rows
        if row["incoming"] or row["outgoing_paid"] or row["outgoing_scheduled"]
    ]
    
    company = CompanyProfile.query.first()
    
    return build_cash_flow_docx(
        company=company,
        period=period,
        active_rows=active_rows,
        total_incoming=sum(row["incoming"] for row in rows),
        total_outgoing_paid=sum(row["outgoing_paid"] for row in rows),
        total_outgoing_scheduled=sum(row["outgoing_scheduled"] for row in rows),
        closing_balance=running_balance,
        show_seal=show_seal,
        show_signatures=show_signatures,
        now=datetime.now(),
    )


@finance_bp.route("/bdds/forecast/print")
@login_required
@finance_required
def bdds_forecast_print():
    """Печать прогноза денежных потоков"""
    start_date_str = request.args.get("start_date", datetime.now().strftime("%Y-%m-%d"))
    show_seal = request.args.get("show_seal", "1") in ["1", "true", "on"]
    show_signatures = request.args.get("show_signatures", "1") in ["1", "true", "on"]

    try:
        start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
    except ValueError:
        start_date = datetime.now()
    start_date = datetime.combine(start_date.date(), datetime.min.time())
    end_date = start_date + timedelta(days=30)
    rows = _build_cash_calendar_rows(start_date, end_date)

    if rows:
        critical_dates = [r["date"] for r in rows if r["balance"] < 0]
    else:
        critical_dates = []

    company = CompanyProfile.query.first()

    return build_bdds_forecast_docx(
        company=company,
        rows=rows,
        ending_balance=rows[-1]["balance"] if rows else 0,
        critical_dates=critical_dates,
        show_seal=show_seal,
        show_signatures=show_signatures,
        now=datetime.now(),
    )


@finance_bp.route("/settlements/print")
@login_required
@finance_required
def settlements_print():
    """Печать расчетов с контрагентами"""
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    show_seal = request.args.get("show_seal", "1") in ["1", "true", "on"]
    show_signatures = request.args.get("show_signatures", "1") in ["1", "true", "on"]
    
    start_date, end_date = get_period_bounds(period)
    
    customer_rows = []
    for customer in Customer.query.order_by(Customer.name.asc()).all():
        orders = [
            order
            for order in customer.sales_orders
            if start_date <= order.order_date < end_date
        ]
        if not orders:
            continue
        billed = sum(order.total_amount for order in orders)
        paid = sum(
            payment.amount
            for order in orders
            for payment in order.payments
            if start_date <= payment.payment_date < end_date
            and payment.status == "completed"
        )
        customer_rows.append(
            {
                "name": customer.name,
                "type": customer.type,
                "orders_count": len(orders),
                "billed": billed,
                "paid": paid,
                "debt": billed - paid,
            }
        )
    
    supplier_rows = []
    for supplier in Supplier.query.order_by(Supplier.name.asc()).all():
        orders = [
            order
            for order in supplier.purchase_orders
            if start_date <= order.order_date < end_date
        ]
        if not orders:
            continue
        accrued = sum(order.total_amount for order in orders)
        paid = sum(order.total_amount for order in orders if order.is_paid)
        supplier_rows.append(
            {
                "name": supplier.name,
                "orders_count": len(orders),
                "accrued": accrued,
                "paid": paid,
                "payable": accrued - paid,
            }
        )
    
    customer_rows.sort(key=lambda row: row["debt"], reverse=True)
    supplier_rows.sort(key=lambda row: row["payable"], reverse=True)
    
    company = CompanyProfile.query.first()
    
    return build_settlements_docx(
        company=company,
        period=period,
        customer_rows=customer_rows,
        supplier_rows=supplier_rows,
        total_customer_billed=sum(row["billed"] for row in customer_rows),
        total_customer_paid=sum(row["paid"] for row in customer_rows),
        total_supplier_accrued=sum(row["accrued"] for row in supplier_rows),
        total_supplier_paid=sum(row["paid"] for row in supplier_rows),
        top_customer_income=sorted(
            calculate_period_customer_income(start_date, end_date).items(),
            key=lambda item: item[1],
            reverse=True,
        )[:5],
        top_supplier_payables=sorted(
            calculate_period_supplier_payables(start_date, end_date).items(),
            key=lambda item: item[1],
            reverse=True,
        )[:5],
        show_seal=show_seal,
        show_signatures=show_signatures,
        now=datetime.now(),
    )


@finance_bp.route("/budget/print")
@login_required
@finance_required
def budget_print():
    """Печать бюджета доходов и расходов"""
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    show_seal = request.args.get("show_seal", "1") in ["1", "true", "on"]
    show_signatures = request.args.get("show_signatures", "1") in ["1", "true", "on"]
    
    items = (
        BudgetItem.query.filter_by(period=period)
        .order_by(BudgetItem.item_type, BudgetItem.category)
        .all()
    )
    income_items = [item for item in items if item.item_type == "income"]
    expense_items = [item for item in items if item.item_type == "expense"]
    total_income_plan = sum(item.planned_amount for item in income_items)
    total_expense_plan = sum(item.planned_amount for item in expense_items)
    
    company = CompanyProfile.query.first()
    
    return build_budget_docx(
        company=company,
        period=period,
        income_items=income_items,
        expense_items=expense_items,
        total_income_plan=total_income_plan,
        total_expense_plan=total_expense_plan,
        planned_balance=total_income_plan - total_expense_plan,
        show_seal=show_seal,
        show_signatures=show_signatures,
        now=datetime.now(),
    )


@finance_bp.route("/bdr/print")
@login_required
@finance_required
def bdr_report_print():
    """Печать анализа рентабельности по сегментам"""
    period = request.args.get("period", datetime.now().strftime("%Y-%m"))
    show_seal = request.args.get("show_seal", "1") in ["1", "true", "on"]
    show_signatures = request.args.get("show_signatures", "1") in ["1", "true", "on"]
    
    start_date, end_date = get_period_bounds(period)
    pivot = calculate_bdr_pivot(start_date, end_date)
    total_revenue = sum(x["revenue"] for x in pivot)
    total_cogs = calculate_cogs_fifo(start_date, end_date)
    total_profit = total_revenue - total_cogs
    
    company = CompanyProfile.query.first()
    
    return build_bdr_docx(
        company=company,
        period=period,
        pivot=pivot,
        total_revenue=total_revenue,
        total_cogs=total_cogs,
        total_profit=total_profit,
        show_seal=show_seal,
        show_signatures=show_signatures,
        now=datetime.now(),
    )
