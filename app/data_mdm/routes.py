from functools import wraps

from datetime import datetime

import io
import json
import re
import urllib.request

from docx import Document
from docx.shared import Inches
from flask import abort, flash, redirect, render_template, request, send_file, url_for
from flask_login import current_user, login_required
from sqlalchemy import and_, func, or_

from app.data_mdm import mdm_bp
from app.models import (
    AuditLog,
    CompanyProfile,
    Customer,
    DataModel,
    DataModelField,
    DuplicateAttempt,
    Employee,
    Product,
    RolePermission,
    Supplier,
    User,
    db,
)


def admin_required(view):
    @wraps(view)
    def wrapper(*args, **kwargs):
        if not current_user.is_authenticated or not (
            current_user.is_admin or current_user.role_admin
        ):
            flash("Раздел управления данными доступен только администратору.", "danger")
            return redirect(url_for("index"))
        return view(*args, **kwargs)

    return wrapper


def mdm_readonly_required(view):
    @wraps(view)
    def wrapper(*args, **kwargs):
        if (
            not current_user.is_authenticated
            or not (
                current_user.is_admin
                or current_user.role_admin
                or current_user.can_view_mdm
                or current_user.can_edit_mdm
            )
        ):
            flash("Доступ к данным MDM ограничен. Обратитесь к администратору.", "danger")
            return redirect(url_for("index"))
        return view(*args, **kwargs)

    return wrapper


def mdm_editor_required(view):
    @wraps(view)
    def wrapper(*args, **kwargs):
        if (
            not current_user.is_authenticated
            or not (
                current_user.is_admin
                or current_user.role_admin
                or current_user.can_edit_mdm
            )
        ):
            flash("Изменение данных MDM разрешено только пользователям с правами редактирования.", "danger")
            return redirect(url_for("mdm.index"))
        return view(*args, **kwargs)

    return wrapper


def _format_audit_value(value):
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y %H:%M")
    return value


def _normalize_phone(phone):
    if not phone:
        return None
    digits = re.sub(r"\D", "", phone)
    if digits.startswith("8") and len(digits) == 11:
        digits = "7" + digits[1:]
    if not digits:
        return None
    return "+" + digits


def _normalize_email(email):
    if not email:
        return None
    normalized = email.strip().lower()
    return normalized if normalized else None


def _normalize_inn(inn):
    if not inn:
        return None
    normalized = re.sub(r"\D", "", inn)
    return normalized if normalized else None


def _validate_email(email):
    if not email:
        return True
    return bool(re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email))


def _validate_phone(phone):
    if not phone:
        return True
    digits = re.sub(r"\D", "", phone)
    return 10 <= len(digits) <= 15


def _validate_inn(inn):
    if not inn:
        return False
    normalized = _normalize_inn(inn)
    return len(normalized) in (10, 12)


def _validate_sku(sku):
    if not sku:
        return False
    return bool(re.match(r"^[A-Za-z0-9\-_.]+$", sku))


def _log_duplicate_attempt(entity, attempted_record, attempted_data, duplicate_fields, reason, source="web"):
    log = DuplicateAttempt(
        entity=entity,
        attempted_record=attempted_record,
        attempted_data=json.dumps(attempted_data, ensure_ascii=False),
        duplicate_fields=", ".join(duplicate_fields),
        source=source,
        reason=reason,
    )
    db.session.add(log)
    db.session.commit()


def _find_potential_duplicates(model, partial_values, exact_values):
    filters = []
    for field, value in partial_values.items():
        if value:
            filters.append(getattr(model, field).ilike(f"%{value}%"))
    for field, value in exact_values.items():
        if value:
            filters.append(getattr(model, field) == value)
    if not filters:
        return []
    return model.query.filter(or_(*filters)).all()


def _entity_duplicate_exists(model, values, exclude_id=None):
    filters = []
    for field, value in values.items():
        if value:
            filters.append(getattr(model, field) == value)
    if not filters:
        return False
    query = model.query.filter(or_(*filters))
    if exclude_id:
        query = query.filter(model.id != exclude_id)
    return query.first() is not None


def _validate_schema_field_name(name):
    return bool(re.match(r'^[A-Za-z_][A-Za-z0-9_]*$', name))


def _get_data_model(entity_key):
    model = DataModel.query.filter_by(key=entity_key).first()
    if not model:
        abort(404)
    return model


def _redirect_to_schema(entity_key, message=None):
    if message:
        flash(message, "warning")
    return redirect(url_for("mdm.schema_fields", entity_key=entity_key))


def _build_mdm_audit_entries():
    entries = []

    entity_specs = [
        (
            "Номенклатура",
            Product.query.all(),
            lambda item: item.name,
            lambda item: item.sku,
            "Создание карточки товара",
            "products",
        ),
        (
            "Клиенты",
            Customer.query.all(),
            lambda item: item.name,
            lambda item: item.type,
            "Создание карточки клиента",
            "customers",
        ),
        (
            "Поставщики",
            Supplier.query.all(),
            lambda item: item.name,
            lambda item: item.inn,
            "Создание карточки поставщика",
            "suppliers",
        ),
        (
            "Сотрудники",
            Employee.query.all(),
            lambda item: item.name,
            lambda item: item.position,
            "Создание карточки сотрудника",
            "employees",
        ),
    ]

    for entity_name, rows, title_getter, meta_getter, action, entity_key in entity_specs:
        for row in rows:
            entries.append(
                {
                    "entity": entity_name,
                    "entity_key": entity_key,
                    "action": action,
                    "record_name": title_getter(row),
                    "record_meta": _format_audit_value(meta_getter(row)),
                    "timestamp": row.created_at,
                    "timestamp_label": row.created_at.strftime("%d.%m.%Y %H:%M")
                    if row.created_at
                    else "н/д",
                    "details": "Запись присутствует в мастер-данных.",
                    "status": "success",
                }
            )

    profile = CompanyProfile.query.first()
    if profile:
        entries.append(
            {
                "entity": "Профиль компании",
                "entity_key": "company_profile",
                "action": "Актуализация профиля организации",
                "record_name": profile.short_name or profile.company_name,
                "record_meta": profile.legal_form,
                "timestamp": profile.updated_at or profile.created_at,
                "timestamp_label": (
                    (profile.updated_at or profile.created_at).strftime("%d.%m.%Y %H:%M")
                    if (profile.updated_at or profile.created_at)
                    else "н/д"
                ),
                "details": "Обновлены реквизиты, контакты или печатные атрибуты.",
                "status": "warning",
            }
        )

    for attempt in DuplicateAttempt.query.order_by(DuplicateAttempt.created_at.desc()).limit(50).all():
        entries.append(
            {
                "entity": "Контроль качества данных",
                "entity_key": "duplicate_attempts",
                "action": "Попытка создания дубликата",
                "record_name": attempt.attempted_record,
                "record_meta": attempt.duplicate_fields,
                "timestamp": attempt.created_at,
                "timestamp_label": attempt.created_at.strftime("%d.%m.%Y %H:%M") if attempt.created_at else "н/д",
                "details": attempt.reason,
                "status": "danger",
            }
        )

    for product in Product.query.all():
        if product.stock and product.stock.last_updated:
            entries.append(
                {
                    "entity": "Остатки",
                    "entity_key": "stock",
                    "action": "Обновление складского остатка",
                    "record_name": product.name,
                    "record_meta": product.sku,
                    "timestamp": product.stock.last_updated,
                    "timestamp_label": product.stock.last_updated.strftime("%d.%m.%Y %H:%M"),
                    "details": (
                        f"На складе: {product.qty_on_hand} ед., "
                        f"в резерве: {product.qty_reserved} ед."
                    ),
                    "status": "info",
                }
            )

    entries.sort(key=lambda item: item["timestamp"] or datetime.min, reverse=True)
    return entries


@mdm_bp.route("/schema")
@login_required
@mdm_readonly_required
def schema_index():
    return redirect(url_for("mdm.index"))


@mdm_bp.route("/schema/<entity_key>")
@login_required
@mdm_readonly_required
def schema_fields(entity_key):
    model = _get_data_model(entity_key)
    fields = model.fields.order_by(DataModelField.order.asc(), DataModelField.label.asc()).all()
    return render_template("data_mdm/schema/fields.html", model=model, fields=fields)


@mdm_bp.route("/schema/<entity_key>/fields/create", methods=["GET", "POST"])
@login_required
@mdm_editor_required
def create_schema_field(entity_key):
    model = _get_data_model(entity_key)
    data_types = [
        ("string", "Строка"),
        ("text", "Текст"),
        ("integer", "Целое число"),
        ("float", "Число"),
        ("date", "Дата"),
        ("boolean", "Флаг"),
    ]
    field_types = [
        ("string", "Текстовое поле"),
        ("textarea", "Текстовая область"),
        ("integer", "Числовое поле"),
        ("float", "Числовое поле"),
        ("date", "Поле даты"),
        ("boolean", "Чекбокс"),
        ("select", "Список"),
        ("radio", "Радио"),
        ("autocomplete", "Автодополнение"),
        ("email", "Email"),
        ("url", "URL"),
    ]
    if request.method == "POST":
        # Verify edit permission before processing any form data
        if not (current_user.is_admin or current_user.role_admin or current_user.can_edit_mdm):
            flash("Изменение данных MDM разрешено только пользователям с правами редактирования.", "danger")
            return redirect(url_for("mdm.schema_index"))
        
        name = request.form.get("name", "").strip()
        label = request.form.get("label", "").strip()
        data_type = request.form.get("data_type", "string")
        field_type = request.form.get("field_type", "string")
        max_length = request.form.get("max_length", None)
        options = request.form.get("options", "").strip() or None
        lookup_source = request.form.get("lookup_source", "").strip() or None
        default_value = request.form.get("default_value", "").strip() or None
        validation_regex = request.form.get("validation_regex", "").strip() or None
        placeholder = request.form.get("placeholder", "").strip() or None
        group = request.form.get("group", "").strip() or None
        readonly = request.form.get("readonly") == "on"
        required = request.form.get("is_required") == "on"
        visible = request.form.get("is_visible") != "off"
        order = request.form.get("order", 0)
        help_text = request.form.get("help_text", "").strip()

        if not name or not label:
            flash("Имя поля и метка обязательны для заполнения.", "danger")
            return redirect(url_for("mdm.create_schema_field", entity_key=entity_key))
        if not _validate_schema_field_name(name):
            flash(
                "Техническое имя поля должно начинаться с буквы или нижнего подчеркивания и содержать только латинские буквы, цифры и подчеркивания.",
                "danger",
            )
            return redirect(url_for("mdm.create_schema_field", entity_key=entity_key))
        if DataModelField.query.filter_by(model_id=model.id, name=name).first():
            flash("Поле с таким техническим именем уже существует для этой сущности.", "danger")
            return redirect(url_for("mdm.create_schema_field", entity_key=entity_key))
        if field_type in ("select", "radio") and not options:
            flash("Для списка или радио-поля необходимо задать опции.", "danger")
            return redirect(url_for("mdm.create_schema_field", entity_key=entity_key))
        try:
            max_length = int(max_length) if max_length else None
            order = int(order)
        except ValueError:
            flash("Длина и порядок поля должны быть целыми числами.", "danger")
            return redirect(url_for("mdm.create_schema_field", entity_key=entity_key))

        db.session.add(
            DataModelField(
                model=model,
                name=name,
                label=label,
                data_type=data_type,
                field_type=field_type,
                max_length=max_length,
                options=options,
                lookup_source=lookup_source,
                default_value=default_value,
                validation_regex=validation_regex,
                placeholder=placeholder,
                group=group,
                readonly=readonly,
                is_required=required,
                is_visible=visible,
                order=order,
                help_text=help_text or None,
            )
        )
        db.session.commit()
        flash("Поле схемы данных создано.", "success")
        return redirect(url_for("mdm.schema_fields", entity_key=entity_key))

    return render_template(
        "data_mdm/schema/field_form.html",
        model=model,
        data_types=data_types,
        field_types=field_types,
        action="create",
        field=None,
    )


@mdm_bp.route("/schema/<entity_key>/fields/<int:field_id>/edit", methods=["GET", "POST"])
@login_required
@mdm_editor_required
def edit_schema_field(entity_key, field_id):
    model = _get_data_model(entity_key)
    field = DataModelField.query.filter_by(model_id=model.id, id=field_id).first_or_404()
    data_types = [
        ("string", "Строка"),
        ("text", "Текст"),
        ("integer", "Целое число"),
        ("float", "Число"),
        ("date", "Дата"),
        ("boolean", "Флаг"),
    ]
    field_types = [
        ("string", "Текстовое поле"),
        ("textarea", "Текстовая область"),
        ("integer", "Числовое поле"),
        ("float", "Числовое поле"),
        ("date", "Поле даты"),
        ("boolean", "Чекбокс"),
        ("select", "Список"),
        ("radio", "Радио"),
        ("autocomplete", "Автодополнение"),
        ("email", "Email"),
        ("url", "URL"),
    ]
    if request.method == "POST":
        # Verify edit permission before processing any form data
        if not (current_user.is_admin or current_user.role_admin or current_user.can_edit_mdm):
            flash("Изменение данных MDM разрешено только пользователям с правами редактирования.", "danger")
            return redirect(url_for("mdm.schema_fields", entity_key=entity_key))
        
        label = request.form.get("label", "").strip()
        data_type = request.form.get("data_type", "string")
        field_type = request.form.get("field_type", "string")
        max_length = request.form.get("max_length", None)
        options = request.form.get("options", "").strip() or None
        lookup_source = request.form.get("lookup_source", "").strip() or None
        default_value = request.form.get("default_value", "").strip() or None
        validation_regex = request.form.get("validation_regex", "").strip() or None
        placeholder = request.form.get("placeholder", "").strip() or None
        group = request.form.get("group", "").strip() or None
        readonly = request.form.get("readonly") == "on"
        required = request.form.get("is_required") == "on"
        visible = request.form.get("is_visible") != "off"
        order = request.form.get("order", 0)
        help_text = request.form.get("help_text", "").strip()

        if not label:
            flash("Метка поля обязательна.", "danger")
            return redirect(url_for("mdm.edit_schema_field", entity_key=entity_key, field_id=field_id))
        if field_type in ("select", "radio") and not options:
            flash("Для списка или радио-поля необходимо задать опции.", "danger")
            return redirect(url_for("mdm.edit_schema_field", entity_key=entity_key, field_id=field_id))
        try:
            field.max_length = int(max_length) if max_length else None
            field.order = int(order)
        except ValueError:
            flash("Длина и порядок поля должны быть целыми числами.", "danger")
            return redirect(url_for("mdm.edit_schema_field", entity_key=entity_key, field_id=field_id))

        field.label = label
        field.data_type = data_type
        field.field_type = field_type
        field.options = options
        field.lookup_source = lookup_source
        field.default_value = default_value
        field.validation_regex = validation_regex
        field.placeholder = placeholder
        field.group = group
        field.readonly = readonly
        field.is_required = required
        field.is_visible = visible
        field.help_text = help_text or None
        db.session.commit()

        flash("Поле схемы данных обновлено.", "success")
        return redirect(url_for("mdm.schema_fields", entity_key=entity_key))

    return render_template(
        "data_mdm/schema/field_form.html",
        model=model,
        field=field,
        data_types=data_types,
        field_types=field_types,
        action="edit",
    )


@mdm_bp.route("/schema/<entity_key>/fields/<int:field_id>/delete", methods=["POST"])
@login_required
@mdm_editor_required
def delete_schema_field(entity_key, field_id):
    model = _get_data_model(entity_key)
    field = DataModelField.query.filter_by(model_id=model.id, id=field_id).first_or_404()
    db.session.delete(field)
    db.session.commit()
    flash("Поле схемы данных удалено.", "info")
    return redirect(url_for("mdm.schema_fields", entity_key=entity_key))


@mdm_bp.route("/api/schema/<entity_key>", methods=["GET"])
@login_required
def get_schema_api(entity_key):
    """
    API endpoint to retrieve schema definition for an entity.
    Used by other modules to dynamically generate forms.
    
    Returns JSON with field definitions:
    {
        "entity_key": "customers",
        "label": "Клиенты",
        "fields": [
            {
                "name": "name",
                "label": "ФИО / наименование",
                "type": "string",
                "required": true,
                "max_length": 255,
                "order": 10
            },
            ...
        ]
    }
    """
    from flask import jsonify
    
    model = _get_data_model(entity_key)
    fields = model.fields.order_by(DataModelField.order.asc()).all()
    
    return jsonify({
        "entity_key": model.key,
        "label": model.label,
        "description": model.description,
        "fields": [
            {
                "name": f.name,
                "label": f.label,
                "type": f.data_type,
                "required": f.is_required,
                "visible": f.is_visible,
                "max_length": f.max_length,
                "help_text": f.help_text,
                "order": f.order,
            }
            for f in fields
        ]
    })


@mdm_bp.route("/api/schema/<entity_key>/visible", methods=["GET"])
@login_required
def get_visible_schema_api(entity_key):
    """
    API endpoint to retrieve only visible fields for an entity.
    Used for form rendering in other modules.
    """
    from flask import jsonify
    
    model = _get_data_model(entity_key)
    fields = [f for f in model.fields.order_by(DataModelField.order.asc()).all() if f.is_visible]
    
    return jsonify({
        "entity_key": model.key,
        "fields": [
            {
                "name": f.name,
                "label": f.label,
                "type": f.data_type,
                "required": f.is_required,
                "max_length": f.max_length,
                "help_text": f.help_text,
            }
            for f in fields
        ]
    })


@mdm_bp.route("/")
@login_required
@mdm_readonly_required
def index():
    products = Product.query.all()
    customers = Customer.query.all()
    suppliers = Supplier.query.all()
    users = User.query.all()
    user_roles = []
    if current_user.is_admin or current_user.role_admin:
        user_roles.append("admin")
    if current_user.can_edit_mdm:
        user_roles.append("mdm-editor")
    elif current_user.can_view_mdm:
        user_roles.append("mdm-viewer")

    return render_template(
        "data_mdm/index.html",
        products_count=len(products),
        active_products_count=sum(1 for product in products if product.is_active),
        low_stock_count=sum(1 for product in products if product.qty_on_hand <= 5),
        customers_count=len(customers),
        legal_customers_count=sum(
            1 for customer in customers if customer.type == "legal_entity"
        ),
        suppliers_count=len(suppliers),
        total_users=len(users),
        active_users=sum(1 for user in users if user.is_active),
        incomplete_customer_contacts=sum(
            1 for customer in customers if not customer.phone or not customer.email
        ),
        incomplete_supplier_profiles=sum(
            1
            for supplier in suppliers
            if not supplier.phone or not supplier.email or not supplier.inn
        ),
        incomplete_employee_contacts=sum(
            1 for user in users if not user.phone or not user.email
        ),
        total_reserved=sum(product.qty_reserved for product in products),
        total_on_hand=sum(product.qty_on_hand for product in products),
        user_roles=user_roles,
    )


@mdm_bp.route("/quality")
@login_required
@mdm_readonly_required
def quality_dashboard():
    products = Product.query.order_by(Product.name.asc()).all()
    customers = Customer.query.order_by(Customer.name.asc()).all()
    suppliers = Supplier.query.order_by(Supplier.name.asc()).all()
    employees = Employee.query.order_by(Employee.name.asc()).all()

    critical_products = [product for product in products if product.qty_on_hand <= 5]
    reserved_products = [product for product in products if product.qty_reserved > 0]
    customers_without_contacts = [
        customer for customer in customers if not customer.phone or not customer.email
    ]
    suppliers_without_profile = [
        supplier
        for supplier in suppliers
        if not supplier.phone or not supplier.email or not supplier.inn
    ]
    employees_without_contacts = [
        employee for employee in employees if not employee.phone or not employee.email
    ]
    inactive_records = {
        "products": sum(1 for product in products if not product.is_active),
        "customers": sum(1 for customer in customers if not customer.is_active),
        "suppliers": sum(1 for supplier in suppliers if not supplier.is_active),
        "employees": sum(1 for employee in employees if not employee.is_active),
    }

    filled_fields = 0
    total_fields = 0
    for customer in customers:
        total_fields += 2
        filled_fields += int(bool(customer.phone)) + int(bool(customer.email))
    for supplier in suppliers:
        total_fields += 3
        filled_fields += (
            int(bool(supplier.phone)) + int(bool(supplier.email)) + int(bool(supplier.inn))
        )
    for employee in employees:
        total_fields += 2
        filled_fields += int(bool(employee.phone)) + int(bool(employee.email))

    completeness_pct = (filled_fields / total_fields * 100) if total_fields else 100
    unit_rows = []
    for unit_name in sorted({product.unit for product in products}):
        unit_products = [product for product in products if product.unit == unit_name]
        unit_rows.append(
            {
                "unit": unit_name,
                "count": len(unit_products),
                "on_hand": sum(product.qty_on_hand for product in unit_products),
                "reserved": sum(product.qty_reserved for product in unit_products),
            }
        )

    return render_template(
        "data_mdm/quality.html",
        completeness_pct=completeness_pct,
        filled_fields=filled_fields,
        total_fields=total_fields,
        inactive_records=inactive_records,
        critical_products=critical_products[:12],
        reserved_products=reserved_products[:12],
        customers_without_contacts=customers_without_contacts[:12],
        suppliers_without_profile=suppliers_without_profile[:12],
        employees_without_contacts=employees_without_contacts[:12],
        unit_rows=unit_rows,
        total_on_hand=sum(product.qty_on_hand for product in products),
        total_reserved=sum(product.qty_reserved for product in products),
        critical_count=len(critical_products),
        reserved_count=len(reserved_products),
    )


@mdm_bp.route("/audit-log")
@login_required
@mdm_readonly_required
def audit_log():
    entity = request.args.get("entity", "").strip()
    q = request.args.get("q", "").strip().lower()

    audit_entries = AuditLog.query.order_by(AuditLog.created_at.desc()).all()
    entries = [
        {
            "timestamp_label": entry.created_at.strftime("%d.%m.%Y %H:%M") if entry.created_at else "н/д",
            "entity": entry.entity,
            "entity_key": entry.entity_key,
            "action": entry.action,
            "record_name": entry.record_name,
            "record_meta": entry.record_meta,
            "details": entry.details,
            "status": entry.status,
            "owner": entry.username or "Система",
        }
        for entry in audit_entries
    ]

    if entity:
        entries = [entry for entry in entries if entry["entity_key"] == entity]
    if q:
        entries = [
            entry
            for entry in entries
            if q in (entry["record_name"] or "").lower()
            or q in (entry["record_meta"] or "").lower()
            or q in (entry["action"] or "").lower()
            or q in (entry["details"] or "").lower()
            or q in (entry["owner"] or "").lower()
        ]

    entity_options = [
        ("products", "Номенклатура"),
        ("customers", "Клиенты"),
        ("suppliers", "Поставщики"),
        ("employees", "Сотрудники"),
        ("users", "Пользователи системы"),
        ("company_profile", "Профиль компании"),
        ("roles", "Роли пользователей"),
        ("duplicate_attempts", "Дубли / конфликтные записи"),
    ]

    return render_template(
        "data_mdm/audit_log.html",
        entries=entries[:150],
        entity=entity,
        q=request.args.get("q", "").strip(),
        entity_options=entity_options,
        total_entries=len(entries),
    )


@mdm_bp.route("/products")
@login_required
@mdm_readonly_required
def products_list():
    page = request.args.get("page", 1, type=int)
    q = request.args.get("q", "").strip()
    active = request.args.get("active", "")
    unit = request.args.get("unit", "")

    query = Product.query
    if q:
        query = query.filter(or_(Product.sku.ilike(f"%{q}%"), Product.name.ilike(f"%{q}%")))
    if active == "active":
        query = query.filter(Product.is_active.is_(True))
    elif active == "inactive":
        query = query.filter(Product.is_active.is_(False))
    if unit:
        query = query.filter(Product.unit == unit)

    products = query.order_by(Product.name.asc()).paginate(page=page, per_page=20)
    all_products = Product.query.all()
    return render_template(
        "data_mdm/products/list.html",
        products=products,
        q=q,
        active=active,
        unit=unit,
        total_products=len(all_products),
        active_products=sum(1 for item in all_products if item.is_active),
        low_stock_products=sum(1 for item in all_products if item.qty_on_hand <= 5),
    )


@mdm_bp.route("/products/create", methods=["GET", "POST"])
@login_required
@mdm_editor_required
def create_product():
    return _redirect_to_schema(
        "products",
        "Добавление новых товарных карточек в МДМ больше не выполняется. Управляйте структурой полей через схему данных.",
    )


@mdm_bp.route("/products/<int:product_id>/edit", methods=["GET", "POST"])
@login_required
@mdm_editor_required
def edit_product(product_id):
    return _redirect_to_schema(
        "products",
        "Редактирование товарных карточек в МДМ больше не выполняется. Управляйте структурой полей через схему данных.",
    )


@mdm_bp.route("/products/<int:product_id>/delete", methods=["POST"])
@login_required
@mdm_editor_required
def delete_product(product_id):
    return _redirect_to_schema(
        "products",
        "Удаление товарных карточек в МДМ больше не выполняется. Используйте операционные модули для управления справочными данными.",
    )


@mdm_bp.route("/products/<int:product_id>/certificate", methods=["GET", "POST"])
@login_required
@mdm_editor_required
def edit_product_certificate(product_id):
    product = Product.query.get_or_404(product_id)

    if request.method == "POST":
        # Verify edit permission before processing any form data
        if not (current_user.is_admin or current_user.role_admin or current_user.can_edit_mdm):
            flash("Изменение данных MDM разрешено только пользователям с правами редактирования.", "danger")
            return redirect(url_for("mdm.products_list"))
        
        certificate_link = request.form.get("certificate_link", "").strip()
        if certificate_link:
            product.certificate_link = certificate_link
            db.session.commit()
            flash("Ссылка на сертификат обновлена.", "success")
        else:
            flash("Ссылка на сертификат не может быть пустой.", "danger")
        return redirect(url_for("mdm.products_list"))

    return render_template("data_mdm/products/certificate.html", product=product)


@mdm_bp.route("/customers")
@login_required
@mdm_readonly_required
def customers_list():
    page = request.args.get("page", 1, type=int)
    q = request.args.get("q", "").strip()
    customer_type = request.args.get("customer_type", "")

    query = Customer.query
    if q:
        query = query.filter(
            or_(
                Customer.name.ilike(f"%{q}%"),
                Customer.phone.ilike(f"%{q}%"),
                Customer.email.ilike(f"%{q}%"),
            )
        )
    if customer_type:
        query = query.filter(Customer.type == customer_type)

    customers = query.order_by(Customer.name.asc()).paginate(page=page, per_page=20)
    return render_template(
        "data_mdm/customers/list.html",
        customers=customers,
        q=q,
        customer_type=customer_type,
        total_customers=Customer.query.count(),
        legal_customers=Customer.query.filter_by(type="legal_entity").count(),
        individual_customers=Customer.query.filter_by(type="individual").count(),
    )


@mdm_bp.route("/customers/create", methods=["GET", "POST"])
@login_required
@mdm_editor_required
def create_customer():
    return _redirect_to_schema(
        "customers",
        "Добавление новых клиентских карточек в МДМ больше не выполняется. Управляйте структурой полей через схему данных.",
    )


@mdm_bp.route("/customers/<int:customer_id>/edit", methods=["GET", "POST"])
@login_required
@mdm_editor_required
def edit_customer(customer_id):
    return _redirect_to_schema(
        "customers",
        "Редактирование клиентских карточек в МДМ больше не выполняется. Управляйте структурой полей через схему данных.",
    )


@mdm_bp.route("/customers/<int:customer_id>/delete", methods=["POST"])
@login_required
@mdm_editor_required
def delete_customer(customer_id):
    return _redirect_to_schema(
        "customers",
        "Удаление клиентских карточек в МДМ больше не выполняется. Используйте операционные модули для управления справочными данными.",
    )


@mdm_bp.route("/suppliers")
@login_required
@mdm_readonly_required
def suppliers_list():
    page = request.args.get("page", 1, type=int)
    q = request.args.get("q", "").strip()

    query = Supplier.query
    if q:
        query = query.filter(
            or_(
                Supplier.name.ilike(f"%{q}%"),
                Supplier.phone.ilike(f"%{q}%"),
                Supplier.email.ilike(f"%{q}%"),
                Supplier.inn.ilike(f"%{q}%"),
            )
        )

    suppliers = query.order_by(Supplier.name.asc()).paginate(page=page, per_page=20)
    return render_template(
        "data_mdm/suppliers/list.html",
        suppliers=suppliers,
        q=q,
        total_suppliers=Supplier.query.count(),
        active_suppliers=Supplier.query.filter_by(is_active=True).count(),
    )


@mdm_bp.route("/suppliers/create", methods=["GET", "POST"])
@login_required
@mdm_editor_required
def create_supplier():
    return _redirect_to_schema(
        "suppliers",
        "Добавление новых карточек поставщиков в МДМ больше не выполняется. Управляйте структурой полей через схему данных.",
    )


@mdm_bp.route("/suppliers/<int:supplier_id>/edit", methods=["GET", "POST"])
@login_required
@mdm_editor_required
def edit_supplier(supplier_id):
    return _redirect_to_schema(
        "suppliers",
        "Редактирование карточек поставщиков в МДМ больше не выполняется. Управляйте структурой полей через схему данных.",
    )


@mdm_bp.route("/suppliers/<int:supplier_id>/delete", methods=["POST"])
@login_required
@mdm_editor_required
def delete_supplier(supplier_id):
    return _redirect_to_schema(
        "suppliers",
        "Удаление карточек поставщиков в МДМ больше не выполняется. Используйте операционные модули для управления справочными данными.",
    )


@mdm_bp.route("/employees")
@login_required
@mdm_readonly_required
def employees_list():
    return redirect(url_for("mdm.users_list"))


@mdm_bp.route("/employees/create", methods=["GET", "POST"])
@login_required
@mdm_editor_required
def create_employee():
    return _redirect_to_schema(
        "employees",
        "Добавление новых карточек сотрудников в МДМ больше не выполняется. Управляйте структурой полей через схему данных.",
    )


@mdm_bp.route("/employees/<int:employee_id>/edit", methods=["GET", "POST"])
@login_required
@mdm_editor_required
def edit_employee(employee_id):
    return _redirect_to_schema(
        "employees",
        "Редактирование карточек сотрудников в МДМ больше не выполняется. Управляйте структурой полей через схему данных.",
    )


@mdm_bp.route("/employees/<int:employee_id>/delete", methods=["POST"])
@login_required
@mdm_editor_required
def delete_employee(employee_id):
    return _redirect_to_schema(
        "employees",
        "Удаление карточек сотрудников в МДМ больше не выполняется. Используйте операционные модули для управления справочными данными.",
    )


@mdm_bp.route("/stock")
@login_required
@mdm_readonly_required
def stock_list():
    page = request.args.get("page", 1, type=int)
    q = request.args.get("q", "").strip()
    stock_state = request.args.get("stock_state", "")

    query = Product.query
    if q:
        query = query.filter(or_(Product.sku.ilike(f"%{q}%"), Product.name.ilike(f"%{q}%")))

    products = query.order_by(Product.name.asc()).all()
    if stock_state == "critical":
        products = [product for product in products if product.qty_on_hand <= 5]
    elif stock_state == "reserved":
        products = [product for product in products if product.qty_reserved > 0]
    elif stock_state == "available":
        products = [product for product in products if product.qty_on_hand - product.qty_reserved > 0]

    start = (page - 1) * 20
    end = start + 20
    page_items = products[start:end]

    class SimplePagination:
        def __init__(self, items, page, per_page):
            self.items = items
            self.page = page
            self.per_page = per_page
            self.total = len(products)
            self.pages = max(1, (self.total + per_page - 1) // per_page)
            self.has_prev = page > 1
            self.has_next = page < self.pages
            self.prev_num = page - 1
            self.next_num = page + 1

        def iter_pages(self):
            return range(1, self.pages + 1)

    pagination = SimplePagination(page_items, page, 20)
    return render_template(
        "data_mdm/stock/list.html",
        products=pagination,
        q=q,
        stock_state=stock_state,
        critical_count=sum(1 for product in Product.query.all() if product.qty_on_hand <= 5),
        reserved_count=sum(1 for product in Product.query.all() if product.qty_reserved > 0),
        total_on_hand=sum(product.qty_on_hand for product in Product.query.all()),
    )


@mdm_bp.route("/stock/report")
@login_required
@mdm_readonly_required
def stock_report():
    q = request.args.get("q", "").strip()
    stock_state = request.args.get("stock_state", "")

    query = Product.query
    if q:
        query = query.filter(or_(Product.sku.ilike(f"%{q}%"), Product.name.ilike(f"%{q}%")))

    products = query.order_by(Product.name.asc()).all()
    if stock_state == "critical":
        products = [product for product in products if product.qty_on_hand <= 5]
    elif stock_state == "reserved":
        products = [product for product in products if product.qty_reserved > 0]
    elif stock_state == "available":
        products = [product for product in products if product.qty_on_hand - product.qty_reserved > 0]

    document = Document()
    logo_url = "https://alaci.kz/wp-content/uploads/2022/03/logo-mebelgrad-e1646908558844.png"
    try:
        image_data = urllib.request.urlopen(logo_url, timeout=10).read()
        header = document.sections[0].header
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = header_paragraph.add_run()
        run.add_picture(io.BytesIO(image_data), width=Inches(1.8))
    except Exception:
        pass

    document.add_heading("Срез складских остатков", level=1)
    filter_text = "Все"
    if stock_state == "critical":
        filter_text = "Критический остаток"
    elif stock_state == "reserved":
        filter_text = "Есть резерв"
    elif stock_state == "available":
        filter_text = "Положительный остаток"

    document.add_paragraph(f"Фильтр: {filter_text}")
    document.add_paragraph(f"Поиск: {q if q else '—'}")
    document.add_paragraph(f"Дата отчета: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Артикул"
    hdr_cells[1].text = "Наименование"
    hdr_cells[2].text = "Ед. изм."
    hdr_cells[3].text = "На складе"
    hdr_cells[4].text = "В резерве"
    hdr_cells[5].text = "Доступно"

    for product in products:
        available = product.qty_on_hand - product.qty_reserved
        row_cells = table.add_row().cells
        row_cells[0].text = product.sku or ""
        row_cells[1].text = product.name or ""
        row_cells[2].text = product.unit or ""
        row_cells[3].text = str(product.qty_on_hand)
        row_cells[4].text = str(product.qty_reserved)
        row_cells[5].text = str(available)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name="stock_report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@mdm_bp.route("/users")
@login_required
@admin_required
def users_list():
    page = request.args.get("page", 1, type=int)
    q = request.args.get("q", "").strip()

    query = User.query
    if q:
        query = query.filter(
            or_(
                User.username.ilike(f"%{q}%"),
                User.email.ilike(f"%{q}%"),
                User.phone.ilike(f"%{q}%"),
            )
        )

    users = query.order_by(User.username.asc()).paginate(page=page, per_page=20)
    return render_template(
        "data_mdm/users/list.html",
        users=users,
        q=q,
        total_users=User.query.count(),
        active_users=User.query.filter_by(is_active=True).count(),
    )


@mdm_bp.route("/users/create", methods=["GET", "POST"])
@login_required
@admin_required
def create_user():
    from app.models import User
    
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        password_repeat = request.form.get("password_repeat", "")
        phone = _normalize_phone(request.form.get("phone", "").strip())

        role_admin = request.form.get("role_admin") == "on"
        role_financier = request.form.get("role_financier") == "on"
        role_warehouse = request.form.get("role_warehouse") == "on"
        role_seller = request.form.get("role_seller") == "on"
        role_count = sum([role_admin, role_financier, role_warehouse, role_seller])

        if not all([username, email, password, password_repeat, phone]):
            flash("Заполните все поля формы и укажите телефон.", "danger")
            return redirect(url_for("mdm.create_user"))

        if role_count != 1:
            flash("Выберите одну роль: админ, финансист, кладовщик или продавец.", "danger")
            return redirect(url_for("mdm.create_user"))

        if password != password_repeat:
            flash("Пароли не совпадают.", "danger")
            return redirect(url_for("mdm.create_user"))

        if not _validate_phone(phone):
            flash("Проверьте корректность формата телефона.", "danger")
            return redirect(url_for("mdm.create_user"))

        if User.query.filter_by(username=username).first():
            flash("Пользователь с таким именем уже существует.", "danger")
            return redirect(url_for("mdm.create_user"))

        if User.query.filter_by(email=email).first():
            flash("Пользователь с таким email уже существует.", "danger")
            return redirect(url_for("mdm.create_user"))

        if User.query.filter_by(phone=phone).first():
            flash("Пользователь с таким телефоном уже существует.", "danger")
            return redirect(url_for("mdm.create_user"))

        user = User(
            username=username,
            email=email,
            phone=phone,
            is_active=request.form.get("is_active") == "on",
            role_admin=role_admin,
            role_financier=role_financier,
            role_warehouse=role_warehouse,
            role_seller=role_seller,
            is_admin=role_admin,
            is_finance=role_financier,
            is_data_admin=role_admin,
            is_data_editor=role_admin,
            is_data_viewer=role_admin,
        )
        user.set_password(password)
        db.session.add(user)
        db.session.commit()

        flash("Пользователь добавлен в систему.", "success")
        return redirect(url_for("mdm.users_list"))

    return render_template("data_mdm/users/create.html")


@mdm_bp.route("/users/<int:user_id>/edit", methods=["GET", "POST"])
@login_required
@admin_required
def edit_user(user_id):
    user = User.query.get_or_404(user_id)

    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        phone = _normalize_phone(request.form.get("phone", user.phone or "").strip())
        password = request.form.get("password", "").strip()
        password_repeat = request.form.get("password_repeat", "").strip()

        role_admin = request.form.get("role_admin") == "on"
        role_financier = request.form.get("role_financier") == "on"
        role_warehouse = request.form.get("role_warehouse") == "on"
        role_seller = request.form.get("role_seller") == "on"
        role_count = sum([role_admin, role_financier, role_warehouse, role_seller])

        if not email or not phone:
            flash("Email и телефон обязательны.", "danger")
            return redirect(url_for("mdm.edit_user", user_id=user_id))

        if role_count != 1:
            flash("Выберите одну роль: админ, финансист, кладовщик или продавец.", "danger")
            return redirect(url_for("mdm.edit_user", user_id=user_id))

        if password and password != password_repeat:
            flash("Пароли не совпадают.", "danger")
            return redirect(url_for("mdm.edit_user", user_id=user_id))

        existing_user = User.query.filter(
            User.email == email,
            User.id != user_id
        ).first()
        if existing_user:
            flash("Пользователь с таким email уже существует.", "danger")
            return redirect(url_for("mdm.edit_user", user_id=user_id))

        existing_phone = User.query.filter(
            User.phone == phone,
            User.id != user_id
        ).first()
        if existing_phone:
            flash("Пользователь с таким телефоном уже существует.", "danger")
            return redirect(url_for("mdm.edit_user", user_id=user_id))

        if not _validate_phone(phone):
            flash("Проверьте корректность формата телефона.", "danger")
            return redirect(url_for("mdm.edit_user", user_id=user_id))

        if password:
            user.set_password(password)

        user.email = email
        user.phone = phone
        user.is_active = request.form.get("is_active") == "on"
        user.role_admin = role_admin
        user.role_financier = role_financier
        user.role_warehouse = role_warehouse
        user.role_seller = role_seller
        user.is_admin = role_admin
        user.is_finance = role_financier
        user.is_data_admin = role_admin
        user.is_data_editor = role_admin
        user.is_data_viewer = role_admin
        db.session.commit()

        flash("Данные пользователя обновлены.", "success")
        return redirect(url_for("mdm.users_list"))

    return render_template("data_mdm/users/edit.html", user=user)


@mdm_bp.route("/users/<int:user_id>/delete", methods=["POST"])
@login_required
@admin_required
def delete_user(user_id):
    from app.models import User
    
    # Prevent self-deletion
    if user_id == current_user.id:
        flash("Вы не можете удалить свой аккаунт.", "danger")
        return redirect(url_for("mdm.users_list"))

    user = User.query.get_or_404(user_id)
    db.session.delete(user)
    db.session.commit()
    flash("Пользователь удален из системы.", "info")
    return redirect(url_for("mdm.users_list"))


@mdm_bp.route("/users/roles", methods=["GET", "POST"])
@login_required
@admin_required
def user_roles():
    role_definitions = [
        {"name": "admin", "label": "Админ"},
        {"name": "finance", "label": "Финансист"},
        {"name": "warehouse", "label": "Кладовщик"},
        {"name": "seller", "label": "Продавец"},
    ]

    def get_role_permission(role_name, label):
        permission = RolePermission.query.filter_by(role_name=role_name).first()
        if permission is None:
            permission = RolePermission(role_name=role_name, label=label)
            db.session.add(permission)
        else:
            permission.label = label
        return permission

    if request.method == "POST":
        for role_definition in role_definitions:
            permission = get_role_permission(
                role_definition["name"], role_definition["label"]
            )
            permission.can_view_mdm = (
                request.form.get(f"can_view_mdm_{role_definition['name']}") == "on"
            )
            permission.can_edit_mdm = (
                request.form.get(f"can_edit_mdm_{role_definition['name']}") == "on"
            )
            permission.can_view_finance = (
                request.form.get(f"can_view_finance_{role_definition['name']}") == "on"
            )
            permission.can_edit_finance = (
                request.form.get(f"can_edit_finance_{role_definition['name']}") == "on"
            )
            permission.can_view_warehouse = (
                request.form.get(f"can_view_warehouse_{role_definition['name']}") == "on"
            )
            permission.can_edit_warehouse = (
                request.form.get(f"can_edit_warehouse_{role_definition['name']}") == "on"
            )
            permission.can_view_sales = (
                request.form.get(f"can_view_sales_{role_definition['name']}") == "on"
            )
            permission.can_edit_sales = (
                request.form.get(f"can_edit_sales_{role_definition['name']}") == "on"
            )

        db.session.commit()
        flash("Права доступа по ролям обновлены", "success")
        return redirect(url_for("mdm.user_roles"))

    roles = [get_role_permission(role["name"], role["label"]) for role in role_definitions]
    db.session.commit()
    return render_template("data_mdm/users/roles.html", roles=roles)


@mdm_bp.route("/company-profile")
@login_required
@mdm_readonly_required
def company_profile():
    profile = CompanyProfile.query.first()
    if not profile:
        flash("Профиль компании не найден. Обратитесь к администратору.", "danger")
        return redirect(url_for("mdm.index"))
    return render_template("data_mdm/company_profile.html", profile=profile)


@mdm_bp.route("/company-profile/edit", methods=["GET", "POST"])
@login_required
@mdm_editor_required
def edit_company_profile():
    profile = CompanyProfile.query.first()
    if not profile:
        flash("Профиль компании не найден. Обратитесь к администратору.", "danger")
        return redirect(url_for("mdm.index"))

    if request.method == "POST":
        # Verify edit permission before processing any form data
        if not (current_user.is_admin or current_user.role_admin or current_user.can_edit_mdm):
            flash("Изменение данных MDM разрешено только пользователям с правами редактирования.", "danger")
            return redirect(url_for("mdm.company_profile"))
        
        # Вспомогательная функция для безопасной обработки формы
        def get_form_value(key, current_value=None, required=False):
            value = request.form.get(key, "").strip()
            if not value:
                return None if not required else (current_value or "")
            return value
        
        profile.company_name = get_form_value("company_name", profile.company_name, required=True)
        profile.short_name = get_form_value("short_name", profile.short_name)
        profile.legal_form = get_form_value("legal_form", profile.legal_form, required=True)
        profile.inn = get_form_value("inn", profile.inn, required=True)
        profile.kpp = get_form_value("kpp", profile.kpp)
        profile.ogrn = get_form_value("ogrn", profile.ogrn, required=True)
        profile.legal_address = get_form_value("legal_address", profile.legal_address, required=True)
        profile.actual_address = get_form_value("actual_address", profile.actual_address)
        profile.phone = get_form_value("phone", profile.phone)
        profile.email = get_form_value("email", profile.email)
        profile.website = get_form_value("website", profile.website)
        profile.bank_name = get_form_value("bank_name", profile.bank_name)
        profile.bank_bik = get_form_value("bank_bik", profile.bank_bik)
        profile.correspondent_account = get_form_value("correspondent_account", profile.correspondent_account)
        profile.settlement_account = get_form_value("settlement_account", profile.settlement_account)
        profile.ceo = get_form_value("ceo", profile.ceo, required=True)
        profile.ceo_position = get_form_value("ceo_position", profile.ceo_position)
        profile.chief_accountant_name = get_form_value("chief_accountant_name", profile.chief_accountant_name)
        profile.logo_url = get_form_value("logo_url", profile.logo_url)
        profile.seal_url = get_form_value("seal_url", profile.seal_url)
        profile.signature_url = get_form_value("signature_url", profile.signature_url)

        # Для ИП не нужны отдельные поля руководителя и главбуха
        if profile.legal_form == "ИП":
            profile.ceo_position = None
            profile.chief_accountant_name = None
            profile.chief_accountant_signature_url = None

        db.session.commit()
        flash("Профиль компании обновлен.", "success")
        return redirect(url_for("mdm.company_profile"))

    return render_template("data_mdm/company_profile_edit.html", profile=profile)


@mdm_bp.route("/backup-settings")
@login_required
@admin_required
def backup_settings():
    from app.models import BackupSettings
    
    settings = BackupSettings.query.first()
    if not settings:
        settings = BackupSettings()
        db.session.add(settings)
        db.session.commit()
    
    return render_template("data_mdm/backup_settings.html", settings=settings)


@mdm_bp.route("/backup-settings", methods=["POST"])
@login_required
@admin_required
def update_backup_settings():
    from app.models import BackupSettings
    from datetime import datetime, timedelta
    
    settings = BackupSettings.query.first()
    if not settings:
        settings = BackupSettings()
        db.session.add(settings)
    
    settings.is_enabled = request.form.get("is_enabled") == "on"
    settings.frequency = request.form.get("frequency", "daily")
    settings.backup_time = request.form.get("backup_time", "02:00")
    settings.retention_days = int(request.form.get("retention_days", 30))
    settings.backup_path = request.form.get("backup_path", "./backups")
    
    # Calculate next backup time
    if settings.is_enabled:
        now = datetime.utcnow()
        if settings.frequency == "daily":
            # Next backup at the specified time today or tomorrow
            today_backup = now.replace(hour=int(settings.backup_time.split(":")[0]), 
                                     minute=int(settings.backup_time.split(":")[1]), 
                                     second=0, microsecond=0)
            if today_backup <= now:
                today_backup += timedelta(days=1)
            settings.next_backup = today_backup
        elif settings.frequency == "weekly":
            # Next Monday at the specified time
            days_until_monday = (7 - now.weekday()) % 7
            if days_until_monday == 0 and now.time() >= datetime.strptime(settings.backup_time, "%H:%M").time():
                days_until_monday = 7
            next_monday = (now + timedelta(days=days_until_monday)).replace(hour=int(settings.backup_time.split(":")[0]), 
                                                                           minute=int(settings.backup_time.split(":")[1]), 
                                                                           second=0, microsecond=0)
            settings.next_backup = next_monday
        elif settings.frequency == "monthly":
            # First day of next month at the specified time
            if now.month == 12:
                next_month = now.replace(year=now.year + 1, month=1, day=1)
            else:
                next_month = now.replace(month=now.month + 1, day=1)
            next_month = next_month.replace(hour=int(settings.backup_time.split(":")[0]), 
                                          minute=int(settings.backup_time.split(":")[1]), 
                                          second=0, microsecond=0)
            settings.next_backup = next_month
    
    db.session.commit()
    flash("Настройки резервного копирования обновлены.", "success")
    return redirect(url_for("mdm.backup_settings"))

