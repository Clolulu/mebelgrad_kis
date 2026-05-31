"""Word-document generation utilities for the warehouse module."""
from __future__ import annotations

import os
import urllib.request
import tempfile
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _cell_borders(cell, top=True, bottom=True, left=True, right=True, color='999999', sz='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, flag in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if flag:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), sz)
            el.set(qn('w:color'), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _para_font(para, bold=False, size=None, color=None, align=None):
    for run in para.runs:
        if bold is not None:
            run.bold = bold
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    if align:
        para.alignment = align


def _add_run(para, text, bold=False, size=10, color=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run


def _format_price(value: float) -> str:
    try:
        text = f'{value:,.2f}'.replace(',', ' ')
        return text.replace('.', ',')
    except Exception:
        return '0,00'


def _download_image(url: str) -> str | None:
    """Download image from URL to a temp file; return path or None."""
    if not url:
        return None
    try:
        if url.startswith('/') or url.startswith('static/'):
            # local file path relative to project root
            base = os.path.join(os.path.dirname(__file__), '..', '..')
            local = os.path.join(base, url.lstrip('/'))
            if os.path.exists(local):
                return local
            return None
        suffix = '.png' if 'png' in url else '.jpg'
        tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
        urllib.request.urlretrieve(url, tmp.name)
        return tmp.name
    except Exception:
        return None


def _sig_block(table, col_idx, title: str, name: str, sig_url: str | None, date_str: str):
    """Fill one column of the 3-column signature table."""
    cell = table.cell(0, col_idx)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

    # title
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, title, bold=True, size=9)

    # signature image or blank line
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sig_url:
        img_path = _download_image(sig_url)
        if img_path:
            try:
                run = p2.add_run()
                run.add_picture(img_path, height=Cm(1.5))
            except Exception:
                _add_run(p2, ' ' * 30, size=10)
        else:
            _add_run(p2, ' ' * 30, size=10)
    else:
        _add_run(p2, ' ' * 30, size=10)

    # underline for name
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run('_' * 28)
    r.font.size = Pt(9)

    # name
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p4, name or '________________________', size=9, color='666666', italic=True)

    # date
    p5 = cell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p5, date_str, size=8, color='888888')


# ── picking doc generator ─────────────────────────────────────────────────────

def generate_picking_docx(pk, company) -> BytesIO:
    """Generate a picking slip .docx for the given PickingOrder."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    order = pk.sales_order
    customer = order.customer if order else None

    # ── Header ────────────────────────────────────────────────────────────────
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_table.columns[0].width = Cm(3.5)
    hdr_table.columns[1].width = Cm(13)

    logo_cell = hdr_table.cell(0, 0)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if company and company.logo_url:
        img_path = _download_image(company.logo_url)
        if img_path:
            try:
                logo_cell.paragraphs[0].add_run().add_picture(img_path, width=Cm(3))
            except Exception:
                pass

    req_cell = hdr_table.cell(0, 1)
    req_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    rp = req_cell.paragraphs[0]
    if company:
        name_line = f'{company.legal_form} «{company.company_name}»' if company.legal_form and company.company_name else (company.company_name or '')
        _add_run(rp, name_line + '\n', bold=True, size=11)
        reqs = []
        if company.inn: reqs.append(f'ИНН {company.inn}')
        if company.kpp: reqs.append(f'КПП {company.kpp}')
        if reqs:
            _add_run(rp, '  '.join(reqs) + '\n', size=9, color='444444')
        if company.legal_address:
            _add_run(rp, f'Адрес: {company.legal_address}\n', size=9, color='444444')
    else:
        _add_run(rp, 'Организация не указана', size=10, color='888888')

    doc.add_paragraph()

    # ── Title ─────────────────────────────────────────────────────────────────
    created_str = pk.created_at.strftime('%d.%m.%Y') if pk.created_at else datetime.utcnow().strftime('%d.%m.%Y')
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(tp, 'ЗАДАНИЕ НА КОМПЛЕКТОВКУ', bold=True, size=14)
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sp, f'№ {pk.picking_number}  от  {created_str}', size=12)
    doc.add_paragraph()

    # ── Order / customer info ─────────────────────────────────────────────────
    info = doc.add_table(rows=6, cols=2)
    info.style = 'Table Grid'
    info.columns[0].width = Cm(5)
    info.columns[1].width = Cm(11.5)

    def _ir(row_idx, label, value):
        lc = info.cell(row_idx, 0)
        _set_cell_bg(lc, 'EBF3FA')
        _add_run(lc.paragraphs[0], label, bold=True, size=9)
        _add_run(info.cell(row_idx, 1).paragraphs[0], value or '—', size=9)

    _ir(0, 'Номер заказа', order.order_number if order else '—')
    _ir(1, 'Дата заказа', order.order_date.strftime('%d.%m.%Y') if order and order.order_date else '—')
    _ir(2, 'Клиент', customer.name if customer else '—')
    _ir(3, 'Телефон клиента', customer.phone if customer else '—')
    _ir(4, 'Адрес доставки', order.delivery_address if order and order.delivery_address else '—')
    _ir(5, 'Комплектовщик', pk.picker.username if pk.picker else '—')

    doc.add_paragraph()

    # ── Items table ───────────────────────────────────────────────────────────
    items_title = doc.add_paragraph()
    _add_run(items_title, 'Состав заказа:', bold=True, size=10)

    items = order.items if order else []
    col_widths2 = [Cm(0.8), Cm(2.2), Cm(6.0), Cm(1.5), Cm(1.5), Cm(2.0), Cm(2.5)]
    headers2 = ['№', 'Артикул', 'Наименование товара', 'Ед.', 'Кол-во', 'На складе', '✓ Выдано']
    tbl = doc.add_table(rows=1 + len(items), cols=7)
    tbl.style = 'Table Grid'

    for ci, hdr in enumerate(headers2):
        cell = tbl.cell(0, ci)
        _set_cell_bg(cell, TABLE_HEADER_COLOR)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, hdr, bold=True, size=9)

    for ri, item in enumerate(items, start=1):
        bg = ALT_ROW_COLOR if ri % 2 == 0 else 'FFFFFF'
        stock_qty = item.product.qty_on_hand if item.product else 0
        ok = stock_qty >= (item.quantity or 0)
        vals2 = [
            (str(ri), WD_ALIGN_PARAGRAPH.CENTER),
            (item.product.sku if item.product else '—', WD_ALIGN_PARAGRAPH.CENTER),
            (item.product.name if item.product else '—', WD_ALIGN_PARAGRAPH.LEFT),
            (item.product.unit if item.product else '—', WD_ALIGN_PARAGRAPH.CENTER),
            (str(item.quantity or 0), WD_ALIGN_PARAGRAPH.CENTER),
            (str(stock_qty), WD_ALIGN_PARAGRAPH.CENTER),
            ('', WD_ALIGN_PARAGRAPH.CENTER),
        ]
        for ci, (val, align) in enumerate(vals2):
            cell2 = tbl.cell(ri, ci)
            _set_cell_bg(cell2, bg)
            p2 = cell2.paragraphs[0]
            p2.alignment = align
            run = _add_run(p2, val, size=9)
            if ci == 5:
                run.font.color.rgb = RGBColor(0x19, 0x87, 0x54) if ok else RGBColor(0xDC, 0x35, 0x45)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Timeline log ──────────────────────────────────────────────────────────
    log_p = doc.add_paragraph()
    _add_run(log_p, 'Журнал выполнения:', bold=True, size=10)
    log_tbl = doc.add_table(rows=4, cols=3)
    log_tbl.style = 'Table Grid'
    log_tbl.columns[0].width = Cm(5)
    log_tbl.columns[1].width = Cm(5)
    log_tbl.columns[2].width = Cm(6.5)

    def _lr(row_idx, label, value):
        lc = log_tbl.cell(row_idx, 0)
        _set_cell_bg(lc, 'EBF3FA')
        _add_run(lc.paragraphs[0], label, bold=True, size=9)
        _add_run(log_tbl.cell(row_idx, 1).paragraphs[0], value or '—', size=9)
        _add_run(log_tbl.cell(row_idx, 2).paragraphs[0], '', size=9)

    _lr(0, 'Задание создано', pk.created_at.strftime('%d.%m.%Y %H:%M') if pk.created_at else '—')
    _lr(1, 'Комплектовка начата', pk.started_at.strftime('%d.%m.%Y %H:%M') if pk.started_at else '—')
    _lr(2, 'Собрано', pk.assembled_at.strftime('%d.%m.%Y %H:%M') if pk.assembled_at else '—')
    _lr(3, 'Отгружено', pk.shipped_at.strftime('%d.%m.%Y %H:%M') if pk.shipped_at else '—')

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Signatures ────────────────────────────────────────────────────────────
    sig_p = doc.add_paragraph()
    _add_run(sig_p, 'Подписи:', bold=True, size=10)
    sig_tbl = doc.add_table(rows=1, cols=3)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for col in sig_tbl.columns:
        col.width = Cm(5.5)

    picker_name = pk.picker.username if pk.picker else '________________________'
    _sig_block(sig_tbl, 0, 'Комплектовал:', picker_name, None,
               pk.assembled_at.strftime('%d.%m.%Y') if pk.assembled_at else '«__» ________ 20__ г.')
    _sig_block(sig_tbl, 1, 'Проверил:', '________________________', None, '«__» ________ 20__ г.')
    _sig_block(sig_tbl, 2, 'Отгрузил:', '________________________', None,
               pk.shipped_at.strftime('%d.%m.%Y') if pk.shipped_at else '«__» ________ 20__ г.')

    if company and company.seal_url:
        img_path = _download_image(company.seal_url)
        if img_path:
            doc.add_paragraph()
            seal_p = doc.add_paragraph()
            try:
                seal_p.add_run().add_picture(img_path, width=Cm(3.5))
            except Exception:
                pass

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── main generator ─────────────────────────────────────────────────────────────

PRIORITY_LABELS = {
    'low': 'Низкий', 'normal': 'Обычный', 'high': 'Высокий', 'urgent': 'Срочный',
}

STATUS_LABELS = {
    'draft': 'Черновик', 'submitted': 'На согласовании',
    'approved': 'Одобрена', 'rejected': 'Отклонена', 'ordered': 'Передана в заказ',
}

TABLE_HEADER_COLOR = 'D6E4F0'
ALT_ROW_COLOR = 'F7FAFD'


def generate_purchase_request_docx(pr, company) -> BytesIO:
    """
    Generate a .docx for the given PurchaseRequest object.
    `company` may be None (CompanyProfile instance or None).
    Returns a BytesIO buffer ready to be sent as file response.
    """
    doc = Document()

    # ── page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    # ── default paragraph style ──────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    # ── HEADER: logo + company requisites ───────────────────────────────────
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_table.columns[0].width = Cm(3.5)
    hdr_table.columns[1].width = Cm(13)

    logo_cell = hdr_table.cell(0, 0)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if company and company.logo_url:
        img_path = _download_image(company.logo_url)
        if img_path:
            try:
                logo_cell.paragraphs[0].add_run().add_picture(img_path, width=Cm(3))
            except Exception:
                pass

    req_cell = hdr_table.cell(0, 1)
    req_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    rp = req_cell.paragraphs[0]

    if company:
        name_line = f'{company.legal_form} «{company.company_name}»' if company.legal_form and company.company_name else (company.company_name or '')
        _add_run(rp, name_line + '\n', bold=True, size=11)
        reqs = []
        if company.inn:
            reqs.append(f'ИНН {company.inn}')
        if company.kpp:
            reqs.append(f'КПП {company.kpp}')
        if company.ogrn:
            reqs.append(f'ОГРН {company.ogrn}')
        if reqs:
            _add_run(rp, '  '.join(reqs) + '\n', size=9, color='444444')
        if company.legal_address:
            _add_run(rp, f'Адрес: {company.legal_address}\n', size=9, color='444444')
        contacts = []
        if company.phone:
            contacts.append(f'Тел.: {company.phone}')
        if company.email:
            contacts.append(f'E-mail: {company.email}')
        if contacts:
            _add_run(rp, '  '.join(contacts) + '\n', size=9, color='444444')
        if company.bank_name:
            _add_run(rp, f'Банк: {company.bank_name}', size=9, color='444444')
            if company.bank_bik:
                _add_run(rp, f'  БИК {company.bank_bik}', size=9, color='444444')
            if company.settlement_account:
                _add_run(rp, f'  р/с {company.settlement_account}', size=9, color='444444')
    else:
        _add_run(rp, 'Организация не указана', size=10, color='888888')

    doc.add_paragraph()  # spacer

    # ── DOCUMENT TITLE ───────────────────────────────────────────────────────
    date_str = pr.request_date.strftime('%d.%m.%Y') if pr.request_date else datetime.utcnow().strftime('%d.%m.%Y')
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title_p, 'ЗАЯВКА НА ЗАКУПКУ', bold=True, size=14)
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(sub_p, f'№ {pr.request_number}  от  {date_str}', bold=False, size=12)

    doc.add_paragraph()

    # ── META INFO TABLE ───────────────────────────────────────────────────────
    meta = doc.add_table(rows=6, cols=2)
    meta.style = 'Table Grid'
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.columns[0].width = Cm(5)
    meta.columns[1].width = Cm(11.5)

    def _meta_row(row_idx, label, value):
        lc = meta.cell(row_idx, 0)
        _set_cell_bg(lc, 'EBF3FA')
        p = lc.paragraphs[0]
        _add_run(p, label, bold=True, size=9)
        vc = meta.cell(row_idx, 1)
        _add_run(vc.paragraphs[0], value or '—', size=9)

    _meta_row(0, 'Инициатор', pr.creator.username if pr.creator else '—')
    _meta_row(1, 'Поставщик', pr.supplier.name if pr.supplier else '—')
    _meta_row(2, 'Основание', pr.comment or '—')
    _meta_row(3, 'Приоритет', PRIORITY_LABELS.get(pr.priority, pr.priority or '—'))
    _meta_row(4, 'Требуемая дата поставки', pr.needed_by_date.strftime('%d.%m.%Y') if pr.needed_by_date else '—')
    _meta_row(5, 'Статус', STATUS_LABELS.get(pr.status, pr.status or '—'))

    doc.add_paragraph()

    # ── ITEMS TABLE ──────────────────────────────────────────────────────────
    items_title = doc.add_paragraph()
    _add_run(items_title, 'Позиции заявки:', bold=True, size=10)

    col_widths = [Cm(0.8), Cm(2.4), Cm(5.8), Cm(1.5), Cm(1.5), Cm(2.2), Cm(2.3)]
    headers = ['№', 'Артикул', 'Наименование товара', 'Ед.', 'Кол-во', 'Цена (оценка)', 'Сумма']
    items_table = doc.add_table(rows=1 + len(pr.items) + 1, cols=7)
    items_table.style = 'Table Grid'

    for i, w in enumerate(col_widths):
        for row in items_table.rows:
            row.cells[i].width = w

    # header row
    for ci, hdr in enumerate(headers):
        cell = items_table.cell(0, ci)
        _set_cell_bg(cell, TABLE_HEADER_COLOR)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, hdr, bold=True, size=9)

    # data rows
    total = 0.0
    for ri, item in enumerate(pr.items, start=1):
        qty = item.quantity or 0
        price = item.estimated_cost or 0.0
        subtotal = qty * price
        total += subtotal
        bg = ALT_ROW_COLOR if ri % 2 == 0 else 'FFFFFF'
        vals = [
            (str(ri), WD_ALIGN_PARAGRAPH.CENTER),
            (item.product.sku if item.product else '—', WD_ALIGN_PARAGRAPH.CENTER),
            (item.product.name if item.product else '—', WD_ALIGN_PARAGRAPH.LEFT),
            (item.product.unit if item.product else '—', WD_ALIGN_PARAGRAPH.CENTER),
            (str(qty), WD_ALIGN_PARAGRAPH.CENTER),
            (_format_price(price) + ' ₽', WD_ALIGN_PARAGRAPH.RIGHT),
            (_format_price(subtotal) + ' ₽', WD_ALIGN_PARAGRAPH.RIGHT),
        ]
        for ci, (val, align) in enumerate(vals):
            cell = items_table.cell(ri, ci)
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = align
            _add_run(p, val, size=9)

    # total row
    total_row_idx = len(pr.items) + 1
    total_left = items_table.cell(total_row_idx, 0)
    total_left.merge(items_table.cell(total_row_idx, 5))
    _set_cell_bg(total_left, 'D6E4F0')
    tp = total_left.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(tp, 'ИТОГО (оценка):', bold=True, size=9)

    total_right = items_table.cell(total_row_idx, 6)
    _set_cell_bg(total_right, 'D6E4F0')
    trp = total_right.paragraphs[0]
    trp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(trp, _format_price(total) + ' ₽', bold=True, size=9)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── SIGNATURE BLOCKS ──────────────────────────────────────────────────────
    sig_title = doc.add_paragraph()
    _add_run(sig_title, 'Подписи:', bold=True, size=10)

    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for col in sig_table.columns:
        col.width = Cm(5.5)

    today = datetime.utcnow().strftime('%d.%m.%Y')
    approved_date = pr.approved_at.strftime('%d.%m.%Y') if pr.approved_at else '«__» ________ 20__ г.'

    # Заказал (creator)
    creator_name = pr.creator.username if pr.creator else '________________________'
    creator_sig = company.ceo_signature_url if (company and pr.creator and pr.creator.is_admin) else None
    _sig_block(sig_table, 0, 'Заказал:', creator_name, creator_sig,
               pr.created_at.strftime('%d.%m.%Y') if pr.created_at else today)

    # Согласовал (approver)
    approver_name = pr.approver.username if pr.approver else '________________________'
    _sig_block(sig_table, 1, 'Согласовал:', approver_name, None, approved_date)

    # Принял (receiver - blank)
    _sig_block(sig_table, 2, 'Принял:', '________________________', None, '«__» ________ 20__ г.')

    # ── SEAL ─────────────────────────────────────────────────────────────────
    if company and company.seal_url:
        img_path = _download_image(company.seal_url)
        if img_path:
            doc.add_paragraph()
            seal_p = doc.add_paragraph()
            try:
                seal_p.add_run().add_picture(img_path, width=Cm(3.5))
            except Exception:
                pass

    # ── FOOTER ───────────────────────────────────────────────────────────────
    if company and company.print_footer:
        doc.add_paragraph()
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(fp, company.print_footer, size=8, color='888888', italic=True)

    # ── serialize to BytesIO ──────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
