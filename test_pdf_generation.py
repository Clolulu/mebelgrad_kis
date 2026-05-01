import os
print('Testing DOCX generation with python-docx...')

try:
    from docx import Document

    doc = Document()
    doc.add_heading('Test DOCX', level=1)
    doc.add_paragraph('This is a test DOCX generated with python-docx.')
    test_path = 'test_docx.docx'
    doc.save(test_path)

    if os.path.exists(test_path):
        print('DOCX generation: SUCCESS')
        os.remove(test_path)
    else:
        print('DOCX generation: FAILED - file not created')
except Exception as e:
    print('DOCX generation error:', e)

try:
    import sys
    sys.path.insert(0, os.path.dirname(__file__))
    from app.finance.routes import create_docx_document

    doc = create_docx_document('Test Report', period='2024-01')
    temp_path = 'test_report.docx'
    doc.save(temp_path)
    print('App DOCX function imported and executed successfully')
    if os.path.exists(temp_path):
        os.remove(temp_path)
except Exception as e:
    print('App DOCX function error:', e)
